import json, os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

with open(os.path.join(os.path.dirname(__file__), "quiz-data.json"), "r", encoding="utf-8") as f:
    questions = json.load(f)

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

t = doc.add_heading("Chinese Medicine Quiz Videos - 20 Questions", 0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Principle: All 3 options look healthy to Western eyes, but only 1 is correct in TCM")
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0xC9, 0xA3, 0x55)
doc.add_paragraph()

for q in questions:
    doc.add_heading(f"Question {q['n']}", 2)
    p = doc.add_paragraph()
    p.add_run("Q (EN): ").bold = True
    p.add_run(q["q_en"])
    p = doc.add_paragraph()
    p.add_run("Q (中文): ").bold = True
    p.add_run(q["q_zh"])
    doc.add_paragraph("Options:")
    for key in ["a", "b", "c"]:
        p = doc.add_paragraph(f"  {q[key]}", style="List Bullet")
    p = doc.add_paragraph()
    p.add_run("Answer: ").bold = True
    p.add_run(q["ans"]).font.color.rgb = RGBColor(0xC9, 0xA3, 0x55)
    p = doc.add_paragraph()
    p.add_run("Why (EN): ").bold = True
    p.add_run(q["why_en"])
    p = doc.add_paragraph()
    p.add_run("原因（中文）: ").bold = True
    p.add_run(q["why_zh"])
    doc.add_paragraph("---")
    doc.add_paragraph()

out = os.path.join(os.path.dirname(__file__), "..", "Quiz_Video_Questions_20.docx")
doc.save(out)
print(f"Saved: {out}")
print(f"Questions: {len(questions)}")
