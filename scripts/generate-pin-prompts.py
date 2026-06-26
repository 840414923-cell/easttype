from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

STYLE_BLOCK = """Style: warm wellness aesthetic, cream background (#FAF6F0),
gold accents (#C9A355), warm brown text (#2A1F14). Clean modern
infographic with subtle Chinese medicine heritage feel. Small
flat food illustrations beside key items. Soft shadows, generous
whitespace, organized in clear sections. The footer URL should be
very small and unobtrusive at the bottom edge. Professional,
saveable, information-rich. All text in English only."""

FOOTER = 'Small footer at very bottom: "Take the free body type quiz -> myeasterntype.com"'

def make_prompt(title, subtitle, sections, footer_note=""):
    parts = []
    parts.append(f'Create a vertical Pinterest infographic (1000x1500px, 2:3 ratio)')
    parts.append(f'about {title.lower()}.')
    parts.append('')
    parts.append(f'Title in elegant serif font: "{title}"')
    if subtitle:
        parts.append(f'Subtitle in smaller text: "{subtitle}"')
    parts.append('')
    for sec in sections:
        parts.append(sec)
        parts.append('')
    parts.append(FOOTER)
    if footer_note:
        parts.append(f'Additional note: {footer_note}')
    parts.append('')
    parts.append(STYLE_BLOCK)
    return '\n'.join(parts)

PROMPTS = []

# ===== SERIES 1: MENSTRUAL CYCLE (4) =====
SERIES_1 = "Series 1: Menstrual Cycle Food Guides"

PROMPTS.append(("Series 1: Menstrual Cycle", 1, "What to Eat the Week Before Your Period", "Chinese Medicine Food Guide for PMS", [
    'Section 1 — heading "EAT MORE TO FEEL BETTER":',
    'Under subheading "Soothe Mood & Tension":',
    '- Mint Tea — moves stuck energy, eases irritability',
    '- Rose Petal Tea — traditional remedy for PMS mood swings',
    '- Chrysanthemum Tea — cools frustration and headaches',
    '',
    'Under subheading "Reduce Bloating":',
    '- Celery — moves digestion, drains water retention',
    '- Daikon Radish — relieves gas and fullness',
    '- Barley Tea — gentle diuretic, reduces puffiness',
    '',
    'Under subheading "Calm Cravings":',
    '- Banana — natural potassium, stabilizes mood',
    '- Dark Chocolate (70%) — small amount eases tension',
    '- Sweet Potato — satisfies sweet cravings steadily',
    '',
    'Section 2 — heading "AVOID TO REDUCE SYMPTOMS":',
    '- Excess Salt — worsens bloating and breast tenderness',
    '- Too Much Coffee — increases anxiety and cramps',
    '- Alcohol — worsens mood swings',
    '- Ice Cold Drinks — worsens cramps',
    '- Refined Sugar — blood sugar crash intensifies emotions',
]))

PROMPTS.append(("Series 1: Menstrual Cycle", 2, "What to Eat During Your Period", "Chinese Medicine Foods for Menstrual Cramps and Comfort", [
    'Section 1 — heading "EAT FOR COMFORT & WARMTH":',
    'Under subheading "Ease Cramps":',
    '- Ginger Tea — warms the uterus, reduces pain',
    '- Brown Sugar Water with Ginger — traditional Chinese remedy for cramps',
    '- Black Sesame Seeds — nourish Blood, ease lower back ache',
    '',
    'Under subheading "Replenish Energy":',
    '- Red Dates (Jujube) — classic Blood-building food',
    '- Longan Fruit — warms and nourishes after blood loss',
    '- Chicken Soup with Ginger — deeply warming and restorative',
    '- Millet Porridge — easy to digest, gentle on the system',
    '',
    'Under subheading "Reduce Flow Discomfort":',
    '- Spinach — iron-rich, replenishes lost minerals',
    '- Dark Cherry — natural anti-inflammatory',
    '',
    'Section 2 — heading "AVOID DURING YOUR PERIOD":',
    '- Ice Cold Drinks and Foods — cause cramping and blood stagnation',
    '- Raw Salads — too hard to digest, drains energy',
    '- Caffeine — worsens cramps and breast tenderness',
    '- Spicy Food — can increase bleeding for some',
    '- Alcohol — thins blood, worsens fatigue',
]))

PROMPTS.append(("Series 3: Week After Period", 3, "What to Eat the Week After Your Period", "Chinese Medicine Foods to Rebuild Blood and Energy", [
    'Section 1 — heading "REBUILD BLOOD & ENERGY":',
    'Under subheading "Nourish Blood":',
    '- Red Dates (Jujube) — the #1 Blood tonic in Chinese medicine',
    '- Goji Berries — supports Liver and Kidney Blood',
    '- Black Sesame Seeds — rich in iron, nourishes Blood',
    '- Spinach — iron and folate for recovery',
    '- Dark Cherries — supports Blood building',
    '',
    'Under subheading "Restore Energy":',
    '- Chicken Soup with Red Dates — classic post-period recovery meal',
    '- Chinese Yam — strengthens Spleen and Kidney together',
    '- Sweet Potato — builds steady energy back up',
    '- Rice Congee — easy to absorb, gentle recovery food',
    '',
    'Under subheading "Hormonal Balance":',
    '- Black Beans — support Kidney energy, hormonal foundation',
    '- Walnuts — nourish Kidney and brain',
    '',
    'Section 2 — heading "AVOID DURING RECOVERY":',
    '- Extreme Dieting — your body needs nourishment now',
    '- Too Much Cold Food — slows recovery and Blood building',
    '- Excessive Coffee — depletes already low energy reserves',
    '- Very Salty Foods — worsens post-period water retention',
]))

PROMPTS.append(("Series 1: Menstrual Cycle", 4, "What to Eat During Ovulation", "Chinese Medicine Foods for Energy and Balance at Ovulation", [
    'Section 1 — heading "SUPPORT YOUR BODY AT OVULATION":',
    'Under subheading "Boost Natural Energy":',
    '- Goji Berries — support reproductive energy',
    '- Black Beans — nourish Kidney, hormonal foundation',
    '- Walnuts — support Kidney essence and brain',
    '- Chinese Yam — strengthens Spleen and Kidney',
    '',
    'Under subheading "Maintain Balance":',
    '- Mung Beans — clear mild heat that peaks mid-cycle',
    '- Mint Tea — keeps Liver energy flowing smoothly',
    '- Celery — light and refreshing, supports fluid balance',
    '- Pear — cooling and moistening',
    '',
    'Under subheading "Light Protein":',
    '- Tofu — gentle plant protein, cooling nature',
    '- Fish (especially salmon) — supports Kidney and Blood',
    '',
    'Section 2 — heading "AVOID AT OVULATION":',
    '- Excess Sugar — destabilizes hormones mid-cycle',
    '- Deep-Fried Foods — generate internal heat and dampness',
    '- Too Much Dairy — can create phlegm and dampness',
    '- Alcohol — disrupts Liver regulation of hormones',
]))

# ===== SERIES 2: DAILY SYMPTOMS (15) =====

PROMPTS.append(("Series 2: Daily Symptoms", 5, "What to Eat When You Can't Sleep", "Chinese Medicine Foods for Insomnia and Better Sleep", [
    'Section 1 — heading "EAT TO SLEEP BETTER":',
    'Under subheading "Calm the Mind":',
    '- Lotus Seed — traditional remedy to calm the Heart and mind',
    '- Lily Bulb — nourishes Heart Yin, reduces vivid dreams',
    '- Warm Milk with Honey — gentle, soothing before bed',
    '',
    'Under subheading "Nourish Blood for Deeper Rest":',
    '- Red Dates (Jujube) — builds Blood that anchors the mind',
    '- Longan Fruit — calms the spirit, aids falling asleep',
    '- Mulberry — nourishes Blood and Yin together',
    '',
    'Under subheading "Warm Evening Drinks":',
    '- Chamomile Tea — relaxes the nervous system',
    '- Rice Congee (small bowl) — easy to digest, won\'t disturb sleep',
    '- Wheat Tea — traditional remedy for night sweats and restlessness',
    '',
    'Section 2 — heading "AVOID BEFORE BED":',
    '- Coffee after noon — stays in your system for 8 hours',
    '- Alcohol — helps you fall asleep but causes 3 AM waking',
    '- Spicy Dinners — generate heat that disturbs sleep',
    '- Large Late Meals — Stomach working hard keeps you awake',
    '- Sugar in the Evening — causes midnight blood sugar crash',
]))

PROMPTS.append(("Series 2: Daily Symptoms", 6, "What to Eat When You're Always Tired", "Chinese Medicine Foods to Rebuild Your Energy", [
    'Section 1 — heading "EAT TO REBUILD ENERGY":',
    'Under subheading "Strengthen Your Core Energy":',
    '- Rice Congee — pre-digested grain, easiest energy source',
    '- Sweet Potato — builds steady, lasting energy',
    '- Chinese Dates (Jujube) — classic Qi tonic',
    '- Chinese Yam — strengthens Spleen and Kidney together',
    '',
    'Under subheading "Deep Nourishment":',
    '- Chicken Broth with Ginger — warming and restorative',
    '- Millet Porridge — the most easily digested grain',
    '- Pumpkin — supports Spleen function',
    '- Goji Berries — supports Liver and Kidney',
    '',
    'Under subheading "Morning Energy":',
    '- Warm Water with Lemon — gentle wake-up for digestion',
    '- Ginger Tea — warms the digestive fire',
    '',
    'Section 2 — heading "AVOID — THESE DRAIN ENERGY":',
    '- Ice Water — extinguishes digestive fire',
    '- Excessive Raw Salads — too hard for a tired Spleen',
    '- Too Much Coffee — borrows energy from tomorrow',
    '- Skipping Breakfast — starves the Spleen at peak time',
    '- Large Late Dinners — body works all night digesting',
]))

PROMPTS.append(("Series 2: Daily Symptoms", 7, "What to Eat When You're Always Cold", "Chinese Medicine Warming Foods for Cold Hands and Feet", [
    'Section 1 — heading "EAT TO WARM YOUR BODY":',
    'Under subheading "Deep Warming Foods":',
    '- Lamb — the strongest warming meat in Chinese medicine',
    '- Ginger (dried is stronger than fresh) — warms from the core',
    '- Cinnamon — warms the Kidney and circulates warmth',
    '- Walnuts — warm and nourish the Kidney',
    '',
    'Under subheading "Warming Grains and Veggies":',
    '- Leeks and Chives — warming, support Kidney Yang',
    '- Pumpkin — warm and sweet, strengthens Spleen',
    '- Sweet Potato — warm nature, builds energy',
    '- Chestnuts — warm the Kidney, strengthen lower back',
    '',
    'Under subheading "Warming Drinks":',
    '- Ginger Tea with Brown Sugar — classic warming combination',
    '- Warm Water — always better than cold',
    '',
    'Section 2 — heading "AVOID — THESE MAKE YOU COLDER":',
    '- Ice Cold Drinks — the #1 enemy of internal warmth',
    '- Watermelon and Cucumber — cooling foods',
    '- Excessive Raw Foods — require heat to process',
    '- Mung Beans — very cooling',
    '- Citrus Fruits in Winter — cooling nature',
]))

PROMPTS.append(("Series 2: Daily Symptoms", 8, "What to Eat When You're Always Bloated", "Chinese Medicine Foods to Reduce Gas and Bloating", [
    'Section 1 — heading "EAT TO REDUCE BLOATING":',
    'Under subheading "Move Stuck Energy":',
    '- Ginger Tea — warms and moves digestion',
    '- Fennel Seeds — classic anti-bloating remedy',
    '- Cardamom — reduces gas and fullness',
    '- Mint Tea — moves Liver energy, eases tension bloating',
    '',
    'Under subheading "Strengthen Digestion":',
    '- Chinese Yam — strengthens Spleen function',
    '- Millet Porridge — easy to digest, reduces burden',
    '- Cooked Carrots — gentle, warming, easy on gut',
    '- White Radish (Daikon) — moves food downward, reduces gas',
    '',
    'Under subheading "Soothing":',
    '- Rice Congee — gentle and restorative',
    '- Papaya — contains enzymes that aid digestion',
    '',
    'Section 2 — heading "AVOID — THESE CAUSE BLOATING":',
    '- Ice Cold Drinks — slow down digestion',
    '- Raw Vegetables — hard for weak digestion to process',
    '- Excessive Dairy — damp-forming, creates gas',
    '- Carbonated Drinks — introduce air into the system',
    '- Chewing Gum — swallows air',
]))

PROMPTS.append(("Series 2: Daily Symptoms", 9, "What to Eat for Brain Fog", "Chinese Medicine Foods to Clear Your Mind", [
    'Section 1 — heading "EAT TO CLEAR YOUR MIND":',
    'Under subheading "Nourish the Brain":',
    '- Walnuts — shaped like a brain, classic Kidney/brain food',
    '- Goji Berries — support Liver and Kidney, improve focus',
    '- Black Sesame Seeds — nourish Brain and Blood',
    '',
    'Under subheading "Strengthen Spleen for Mental Clarity":',
    '- Rice Congee — easy energy for the mind',
    '- Sweet Potato — builds the energy the brain runs on',
    '- Chinese Dates — nourish Blood that feeds the mind',
    '- Chinese Yam — strengthens Spleen and Kidney',
    '',
    'Under subheading "Move Stagnant Energy":',
    '- Green Tea — gentle caffeine, clears mental fog',
    '- Mint Tea — moves energy, lifts the mind',
    '- Rose Tea — moves Liver Qi, reduces mental tension',
    '',
    'Section 2 — heading "AVOID — THESE WORSEN BRAIN FOG":',
    '- Excessive Sugar — crash creates deeper fog',
    '- Too Much Coffee — temporary clarity then worse crash',
    '- Heavy Greasy Meals — body diverts energy to digestion',
    '- Ice Cold Drinks — slow Spleen function',
    '- Skipping Meals — brain runs out of fuel',
]))

PROMPTS.append(("Series 2: Daily Symptoms", 10, "What to Eat When You Feel Anxious", "Chinese Medicine Foods to Calm Anxiety Naturally", [
    'Section 1 — heading "EAT TO CALM YOUR MIND":',
    'Under subheading "Calm the Heart and Spirit":',
    '- Lotus Seed — traditional remedy to settle the mind',
    '- Lily Bulb — nourishes Heart Yin, reduces restlessness',
    '- Red Dates — builds Blood that grounds the mind',
    '- Wheat — calms the spirit, reduces anxiety',
    '',
    'Under subheading "Move Stuck Liver Energy":',
    '- Chrysanthemum Tea — cools frustration, calms',
    '- Rose Petal Tea — moves Liver Qi, lifts mood',
    '- Mint Tea — light and moving, reduces tension',
    '- Celery — moves energy downward',
    '',
    'Under subheading "Ground and Nourish":',
    '- Warm Rice Congee — grounding and easy',
    '- Bone Broth — deeply nourishing, stabilizes energy',
    '',
    'Section 2 — heading "AVOID — THESE INCREASE ANXIETY":',
    '- Coffee and Energy Drinks — directly stimulate Heart',
    '- Alcohol — temporary relief then worse rebound',
    '- Very Spicy Foods — add internal heat to the Heart',
    '- Skipping Meals — destabilizes Blood Sugar and mood',
    '- Late Night Screens — disturb the spirit',
]))

PROMPTS.append(("Series 2: Daily Symptoms", 11, "What to Eat for Night Sweats", "Chinese Medicine Cooling Foods to Reduce Night Sweats", [
    'Section 1 — heading "EAT TO COOL AND MOISTEN":',
    'Under subheading "Nourish Kidney Yin":',
    '- Black Sesame Seeds — moisten and cool, classic Yin food',
    '- Goji Berries — support Liver and Kidney Yin',
    '- Mulberry — cooling, nourishes Blood and Yin',
    '- Chinese Yam — gently nourishes Kidney',
    '',
    'Under subheading "Cool Internal Heat":',
    '- Mung Beans — strongest cooling food',
    '- Pear — moistening and cooling',
    '- Lotus Root — clears heat',
    '- Watermelon — cooling and hydrating',
    '',
    'Under subheading "Calm the Heart for Sleep":',
    '- Lotus Seed — calms mind, reduces sleep sweating',
    '- Lily Bulb — cools Heart, promotes restful sleep',
    '- Wheat Tea — traditional night sweat remedy',
    '',
    'Section 2 — heading "AVOID — THESE WORSEN NIGHT SWEATS":',
    '- Spicy Foods — generate internal heat',
    '- Alcohol — produces heat and disrupts sleep',
    '- Ginger in Excess — warming, can worsen at night',
    '- Lamb and Cinnamon — too warming for this pattern',
    '- Late Heavy Dinners — generate heat during sleep',
]))

PROMPTS.append(("Series 2: Daily Symptoms", 12, "What to Eat for Acid Reflux", "Chinese Medicine Foods to Cool Stomach Heat", [
    'Section 1 — heading "EAT TO CALM STOMACH HEAT":',
    'Under subheading "Cool and Soothe":',
    '- Mung Bean Soup — strongest cooling food for Stomach',
    '- Pear — cooling and moistening',
    '- Lotus Root — clears heat, soothing',
    '- Cucumber — cooling and hydrating',
    '',
    'Under subheading "Gentle and Easy":',
    '- Millet Porridge — easy on the stomach',
    '- Banana — gentle, coats the stomach',
    '- Chinese Yam — supports Stomach without adding heat',
    '- Papaya — enzymes aid gentle digestion',
    '',
    'Under subheading "Calm and Move":',
    '- Rice Congee — gentle, won\'t trigger reflux',
    '- Cooked Apple — soft, easy, non-acidic',
    '',
    'Section 2 — heading "AVOID — THESE TRIGGER REFLUX":',
    '- Spicy Food — adds fuel to Stomach fire',
    '- Coffee — relaxes the valve, adds heat',
    '- Alcohol — generates Stomach heat',
    '- Fried and Greasy Foods — heavy and heat-forming',
    '- Eating Within 3 Hours of Bed — reflux when lying down',
]))

PROMPTS.append(("Series 2: Daily Symptoms", 13, "What to Eat for Constipation", "Chinese Medicine Foods to Restore Regularity", [
    'Section 1 — heading "EAT FOR SMOOTH ELIMINATION":',
    'Under subheading "Moisten the Intestines":',
    '- Black Sesame Seeds — natural oils lubricate',
    '- Honey in Warm Water — moistens and gently moves',
    '- Pear — cooling and moistening',
    '- Sweet Almond — lubricates intestines',
    '',
    'Under subheading "Move Energy Downward":',
    '- White Radish (Daikon) — moves food downward',
    '- Spinach — cooling, moistening, gentle',
    '- Banana (ripe) — soft and moistening',
    '',
    'Under subheading "Strengthen the Push":',
    '- Sweet Potato — softens stool, strengthens Spleen',
    '- Walnuts — moisten intestines, warm Kidney',
    '- Seaweed — softens, adds minerals and fiber',
    '',
    'Section 2 — heading "AVOID — THESE CAUSE CONSTIPATION":',
    '- Very Spicy Foods — dry out the intestines',
    '- Dry Roasted Nuts without Water — drying',
    '- Persimmon (unripe) — astringent, binding',
    '- Excessive White Flour — dry, no fiber',
    '- Too Much Coffee — depletes body fluids',
]))

PROMPTS.append(("Series 2: Daily Symptoms", 14, "What to Eat for Weight Loss", "Chinese Medicine Foods to Support Healthy Metabolism", [
    'Section 1 — heading "EAT TO SUPPORT HEALTHY WEIGHT":',
    'Under subheading "Drain Dampness":',
    '- Adzuki Beans — strongest damp-draining food',
    '- Coix Seed (Job\'s Tears) — drains damp, strengthens Spleen',
    '- Winter Melon — promotes fluid metabolism',
    '- Lotus Leaf Tea — traditional weight support tea',
    '',
    'Under subheading "Move and Circulate":',
    '- Green Tea — moves Qi, supports metabolism',
    '- Hawthorn — helps digest fats',
    '- Radish (Daikon) — moves Qi, aids digestion',
    '- Barley Tea — drains dampness, light',
    '',
    'Under subheading "Strengthen the Spleen":',
    '- Chinese Yam — supports Spleen function',
    '- Millet — light grain, easy to process',
    '- Cooked Vegetables — easier on digestion than raw',
    '',
    'Section 2 — heading "AVOID — THESE HOLD WEIGHT":',
    '- Dairy Products — most damp-forming food group',
    '- Refined Sugar — feeds dampness and phlegm',
    '- Cold Drinks and Ice Water — slow Spleen function',
    '- Deep-Fried Foods — generate phlegm',
    '- Excessive Raw Foods — hard for Spleen to process',
]))

PROMPTS.append(("Series 2: Daily Symptoms", 15, "What to Eat for Clear Skin", "Chinese Medicine Foods for Acne-Free Skin from Within", [
    'Section 1 — heading "EAT FOR CLEAR SKIN":',
    'Under subheading "Clear Internal Heat":',
    '- Mung Beans — clear heat and dampness',
    '- Green Tea — cools heat, rich in antioxidants',
    '- Cucumber — cooling and hydrating for skin',
    '- Bitter Gourd — clears heat, supports liver',
    '',
    'Under subheading "Nourish Blood for Glow":',
    '- Goji Berries — support Liver Blood for skin',
    '- Black Sesame Seeds — nourish Blood and moisturize',
    '- Red Dates — build Blood, improve complexion',
    '- Mulberry — nourishes Blood and Yin',
    '',
    'Under subheading "Reduce Inflammation":',
    '- Lotus Root — clears heat, improves circulation',
    '- Pear — cooling, moistens skin from within',
    '- Winter Melon — clears heat, reduces oiliness',
    '',
    'Section 2 — heading "AVOID — THESE CAUSE BREAKOUTS":',
    '- Spicy Foods — add heat that shows on face',
    '- Fried and Greasy Foods — create damp-heat',
    '- Excessive Dairy — damp-forming, triggers acne',
    '- Refined Sugar — feeds inflammation',
    '- Chocolate (excess) — can trigger breakouts',
]))

PROMPTS.append(("Series 2: Daily Symptoms", 16, "What to Eat for Hot Flashes", "Chinese Medicine Cooling Foods to Reduce Internal Heat", [
    'Section 1 — heading "EAT TO COOL DOWN":',
    'Under subheading "Strongest Cooling Foods":',
    '- Mung Beans — clear empty heat',
    '- Cucumber — cooling and hydrating',
    '- Watermelon — cools heat quickly',
    '- Pear — moistens and cools',
    '',
    'Under subheading "Nourish Kidney Yin":',
    '- Black Sesame Seeds — rebuild cooling reserve',
    '- Tofu and Soy — cooling, gentle phytoestrogens',
    '- Mulberry — nourishes Blood and Yin',
    '- Lotus Root — clears heat from the body',
    '',
    'Under subheading "Calm and Cool Drinks":',
    '- Mint Tea — cooling, moves energy',
    '- Chrysanthemum Tea — cools Liver heat',
    '- Room Temperature Water — never ice cold',
    '',
    'Section 2 — heading "AVOID — THESE FEED HOT FLASHES":',
    '- Spicy Foods — add fuel to internal fire',
    '- Alcohol — generates heat rapidly',
    '- Excess Coffee — stimulates Heart, adds heat',
    '- Lamb and Cinnamon — too warming',
    '- Saunas and Hot Baths — external heat compounds internal',
]))

PROMPTS.append(("Series 2: Daily Symptoms", 17, "What to Eat for Water Retention", "Chinese Medicine Foods to Reduce Puffiness and Swelling", [
    'Section 1 — heading "EAT TO DRAIN EXCESS FLUID":',
    'Under subheading "Drain Dampness":',
    '- Adzuki Beans — strongest damp-draining food',
    '- Coix Seed (Job\'s Tears) — drains fluid, strengthens Spleen',
    '- Winter Melon — promotes urination, reduces swelling',
    '- Lotus Leaf Tea — reduces water weight',
    '',
    'Under subheading "Support Fluid Metabolism":',
    '- Celery — natural diuretic, moves fluid',
    '- Barley Tea — gentle, drains dampness daily',
    '- Corn Silk Tea — traditional remedy for edema',
    '- Chinese Yam — strengthens Spleen to process fluids',
    '',
    'Under subheading "Move Energy":',
    '- White Radish — moves energy downward',
    '- Green Tea — light, moves Qi',
    '',
    'Section 2 — heading "AVOID — THESE WORSEN RETENTION":',
    '- Excess Salt — holds water in tissues',
    '- Cold Drinks — slow fluid metabolism',
    '- Dairy Products — damp-forming, increases swelling',
    '- Deep-Fried Foods — generate dampness',
    '- Sedentary Habits — lack of movement worsens pooling',
]))

PROMPTS.append(("Series 2: Daily Symptoms", 18, "What to Eat When You Crave Sweets", "Chinese Medicine Foods to Balance Sugar Cravings", [
    'Section 1 — heading "EAT TO CALM SWEET CRAVINGS":',
    'Under subheading "Natural Sweets That Satisfy":',
    '- Sweet Potato — naturally sweet, builds energy steadily',
    '- Chinese Dates (Jujube) — sweet Qi tonic, satisfies',
    '- Goji Berries — mildly sweet, nourishing',
    '- Carrots (cooked) — natural sweetness, gentle',
    '',
    'Under subheading "Stabilize Blood Sugar":',
    '- Millet Porridge — complex carbs, steady energy',
    '- Rice Congee — gentle, prevents sugar crashes',
    '- Chinese Yam — stabilizes and strengthens',
    '- Pumpkin — warm, sweet, supports Spleen',
    '',
    'Under subheading "Reduce the Urge":',
    '- Warm Water with Lemon — reduces false hunger',
    '- Nuts (small amount) — protein stabilizes sugar',
    '- Cooked Apple — natural sweetness without spike',
    '',
    'Section 2 — heading "AVOID — THESE TRIGGER MORE CRAVINGS":',
    '- Refined Sugar — spike and crash cycle',
    '- Pastries and White Flour — rapid blood sugar swing',
    '- Skipping Meals — leads to desperate sugar craving',
    '- Too Much Coffee — destabilizes blood sugar',
    '- Artificial Sweeteners — perpetuate sweet dependence',
]))

PROMPTS.append(("Series 2: Daily Symptoms", 19, "What to Eat for Low Energy Mornings", "Chinese Medicine Foods to Start Your Day with Energy", [
    'Section 1 — heading "EAT FOR MORNING ENERGY":',
    'Under subheading "Warm the Digestive Fire":',
    '- Warm Water with Lemon — gentle wake-up for Spleen',
    '- Ginger Tea — warms digestive fire, starts metabolism',
    '- Rice Congee — easy morning energy, no burden',
    '',
    'Under subheading "Build Steady Energy":',
    '- Sweet Potato — slow-burning energy source',
    '- Chinese Dates (Jujube) — natural morning Qi boost',
    '- Millet Porridge — the best breakfast grain',
    '- Chinese Yam — strengthens Spleen for the day',
    '',
    'Under subheading "Nourish Deeply":',
    '- Chicken Broth — warm and restorative',
    '- Goji Berries in warm water — supports Kidney and Liver',
    '- Walnuts — brain food for mental energy',
    '',
    'Section 2 — heading "AVOID IN THE MORNING":',
    '- Ice Cold Drinks — shock and slow the Spleen',
    '- Coffee on Empty Stomach — burns through borrowed energy',
    '- Skipping Breakfast — starves the Spleen at peak time (7-9 AM)',
    '- Raw Cold Salads — too hard to process in the morning',
    '- Sugary Pastries — energy crash by mid-morning',
]))

# ===== SERIES 3: SEASONAL (4) =====

PROMPTS.append(("Series 3: Seasonal Health", 20, "Cooling Foods for Summer Heat", "Chinese Medicine Guide to Stay Cool from Within", [
    'Section 1 — heading "EAT TO COOL DOWN IN SUMMER":',
    'Under subheading "Strongest Cooling Foods":',
    '- Mung Beans — clear summer heat',
    '- Watermelon — nature\'s summer coolant',
    '- Cucumber — hydrating and cooling',
    '- Bitter Gourd — clears heat strongly',
    '',
    'Under subheading "Cooling Fruits":',
    '- Pear — moistens and cools',
    '- Lotus Root — refreshing, clears heat',
    '- Mulberry — cooling and nourishing',
    '- Tomato — mildly cooling, rich in lycopene',
    '',
    'Under subheading "Cooling Drinks":',
    '- Green Tea — light and cooling',
    '- Mint Tea — refreshing, moves energy',
    '- Chrysanthemum Tea — cools Liver and head',
    '- Barley Tea — cooling, replaces water',
    '',
    'Section 2 — heading "LIMIT IN SUMMER":',
    '- Lamb and Venison — too warming',
    '- Excessive Spicy Food — adds internal heat',
    '- Heavy Greasy Meals — hard to digest in heat',
    '- Too Much Coffee — heating and dehydrating',
    '- Alcohol — generates internal heat',
]))

PROMPTS.append(("Series 3: Seasonal Health", 21, "Warming Foods for Winter Cold", "Chinese Medicine Guide to Stay Warm from Within", [
    'Section 1 — heading "EAT TO WARM UP IN WINTER":',
    'Under subheading "Strongest Warming Foods":',
    '- Lamb — the most warming meat',
    '- Ginger (dried) — deep internal warmth',
    '- Cinnamon — warms Kidney and circulation',
    '- Deer Meat (venison) — very warming',
    '',
    'Under subheading "Warming Vegetables":',
    '- Leeks and Chives — warm the body',
    '- Pumpkin — warm and nourishing',
    '- Sweet Potato — warm energy builder',
    '- Fennel — warming, aids digestion',
    '',
    'Under subheading "Warming Drinks and Grains":',
    '- Ginger Tea with Brown Sugar — classic winter drink',
    '- Walnut and Red Date Tea — warming tonic',
    '- Millet Porridge — warm grain breakfast',
    '- Chestnut Soup — warms Kidney, strengthens',
    '',
    'Section 2 — heading "LIMIT IN WINTER":',
    '- Ice Cold Drinks — destroy winter warmth',
    '- Watermelon — too cooling for cold season',
    '- Excessive Raw Salads — hard on cold digestion',
    '- Mung Beans — very cooling',
    '- Citrus in Excess — can be cooling',
]))

PROMPTS.append(("Series 3: Seasonal Health", 22, "Foods for Spring Allergies", "Chinese Medicine Guide to Reduce Seasonal Allergy Symptoms", [
    'Section 1 — heading "EAT TO REDUCE ALLERGY SYMPTOMS":',
    'Under subheading "Strengthen Lung Wei Qi":',
    '- Chinese Dates — build defensive energy',
    '- Chinese Yam — strengthens Lung and Spleen',
    '- Lotus Seed — supports the immune foundation',
    '- White Fungus — moistens Lung, reduces dryness',
    '',
    'Under subheading "Clear Mild Heat":',
    '- Chrysanthemum Tea — clears heat from head and eyes',
    '- Mint Tea — opens sinuses, moves energy',
    '- Green Tea — gentle heat-clearing',
    '- Mung Bean Soup — mild cooling for allergy season',
    '',
    'Under subheading "Reduce Phlegm":',
    '- Pear — moistens Lung, reduces dry cough',
    '- White Radish — clears phlegm downward',
    '- Lily Bulb — moistens Lung, calms',
    '',
    'Section 2 — heading "AVOID DURING ALLERGY SEASON":',
    '- Dairy Products — increase phlegm and congestion',
    '- Excessive Sugar — feeds inflammation',
    '- Cold Drinks — weaken Lung function',
    '- Deep-Fried Foods — create phlegm',
    '- Ice Cream — double phlegm-forming',
]))

PROMPTS.append(("Series 3: Seasonal Health", 23, "Foods for Dry Autumn Skin", "Chinese Medicine Guide to Moisturize from Within", [
    'Section 1 — heading "EAT TO MOISTEN IN AUTUMN":',
    'Under subheading "Moisten the Lung and Skin":',
    '- Pear — the #1 autumn moisturizing food',
    '- White Fungus — deeply moistening',
    '- Lily Bulb — moistens Lung and skin',
    '- Lotus Root — moistening and nourishing',
    '',
    'Under subheading "Nourish Body Fluids":',
    '- Honey in Warm Water — natural moisturizer',
    '- Sesame Seeds (black) — moisten and nourish',
    '- Sweet Almond — lubricates and moistens',
    '- Mulberry — nourishes Blood and fluids',
    '',
    'Under subheading "Warm and Soothing":',
    '- Rice Congee — easy, moistening meal',
    '- Snow Fungus Soup with Rock Sugar — classic autumn dessert',
    '- Apple (cooked) — warm and moistening',
    '',
    'Section 2 — heading "AVOID IN DRY AUTUMN":',
    '- Very Spicy Foods — dry out the body',
    '- Excessive Coffee — depletes fluids',
    '- Dry Roasted Snacks without hydration',
    '- Too Much Salt — dries out tissues',
    '- Deep-Fried Foods — drying nature',
]))

# ===== SERIES 4: SKIN & BEAUTY (4) =====

PROMPTS.append(("Series 4: Skin & Beauty", 24, "Foods for Clear Skin", "Chinese Medicine Guide to Clear Acne from Within", [
    'Section 1 — heading "EAT FOR CLEAR SKIN":',
    'Under subheading "Clear Internal Heat":',
    '- Mung Beans — clear heat and toxins',
    '- Green Tea — antioxidant and cooling',
    '- Cucumber — hydrating, cooling',
    '- Bitter Gourd — clears liver heat',
    '',
    'Under subheading "Nourish Blood for Glow":',
    '- Goji Berries — improve skin from within',
    '- Black Sesame Seeds — nourish and moisturize',
    '- Red Dates — build Blood for complexion',
    '- Mulberry — rich in antioxidants',
    '',
    'Under subheading "Reduce Inflammation":',
    '- Lotus Root — improves circulation to skin',
    '- Winter Melon — reduces oiliness',
    '- Pear — moistens without adding heat',
    '',
    'Section 2 — heading "AVOID FOR CLEAR SKIN":',
    '- Spicy Foods — add heat to face',
    '- Fried Foods — create damp-heat in skin',
    '- Dairy — triggers acne for many',
    '- Refined Sugar — feeds skin inflammation',
    '- Excess Chocolate — can trigger breakouts',
]))

PROMPTS.append(("Series 4: Skin & Beauty", 25, "Foods for Glowing Complexion", "Chinese Medicine Foods for Radiant Skin from Within", [
    'Section 1 — heading "EAT FOR A GLOWING COMPLEXION":',
    'Under subheading "Build Blood for Radiance":',
    '- Red Dates — the classic beauty food',
    '- Goji Berries — support Liver Blood',
    '- Black Sesame Seeds — nourish and moisturize',
    '- Longan Fruit — sweet Blood tonic',
    '',
    'Under subheading "Improve Circulation":',
    '- Saffron (small amount) — moves Blood beautifully',
    '- Rose Petal Tea — moves Liver Qi, improves color',
    '- Black Beans — nourish Kidney and Blood',
    '- Dark Cherries — rich in skin antioxidants',
    '',
    'Under subheading "Moisten and Hydrate":',
    '- White Fungus — natural skin moisturizer from within',
    '- Pear — hydrating and cooling',
    '- Honey in Warm Water — nourishes and moistens',
    '- Lotus Seed — calms, supports skin repair',
    '',
    'Section 2 — heading "AVOID FOR RADIANT SKIN":',
    '- Excessive Cold Drinks — slow nutrient delivery',
    '- Too Much Sugar — damages collagen',
    '- Greasy Fast Food — dulls complexion',
    '- Excess Coffee — dehydrates skin',
    '- Smoking and Alcohol — age skin rapidly',
]))

PROMPTS.append(("Series 4: Skin & Beauty", 26, "Foods for Hair Health", "Chinese Medicine Foods to Strengthen Hair from Within", [
    'Section 1 — heading "EAT FOR STRONG HAIR":',
    'Under subheading "Nourish Kidney for Hair Root":',
    '- Black Sesame Seeds — the #1 hair food in Chinese medicine',
    '- Black Beans — nourish Kidney, support hair pigment',
    '- Walnuts — warm Kidney, strengthen hair',
    '- Mulberry — nourishes Blood and Kidney',
    '',
    'Under subheading "Build Blood for Hair Shine":',
    '- Red Dates — build Blood that feeds hair',
    '- Goji Berries — support Liver Blood',
    '- Dark Leafy Greens — iron and minerals',
    '- Dried Apricots — Blood nourishing',
    '',
    'Under subheading "Support Scalp Circulation":',
    '- Black Fungus (Wood Ear) — improves circulation',
    '- Ginger Tea — warms and circulates to scalp',
    '- Chrysanthemum Tea — clears head heat',
    '',
    'Section 2 — heading "AVOID FOR HAIR HEALTH":',
    '- Excessive Sugar — damages hair follicles',
    '- Too Much Salt — can weaken Kidney',
    '- Cold Drinks — slow circulation to scalp',
    '- Fried Foods — create damp-heat affecting scalp',
    '- Stress and Late Nights — deplete Kidney essence',
]))

PROMPTS.append(("Series 4: Skin & Beauty", 27, "Foods to Reduce Puffiness", "Chinese Medicine Foods to Depuff Your Face and Body", [
    'Section 1 — heading "EAT TO REDUCE PUFFINESS":',
    'Under subheading "Drain Dampness":',
    '- Adzuki Beans — strongest depuffing food',
    '- Coix Seed (Job\'s Tears) — drain excess fluid',
    '- Winter Melon — natural diuretic',
    '- Lotus Leaf Tea — traditional depuffing tea',
    '',
    'Under subheading "Move Fluids":',
    '- Celery — natural diuretic, moves water',
    '- Corn Silk Tea — reduces water retention',
    '- Barley Tea — daily gentle draining',
    '- Cucumber — cooling, reduces face puffiness',
    '',
    'Under subheading "Strengthen Spleen":',
    '- Chinese Yam — helps Spleen process fluids',
    '- Millet — light, supports fluid metabolism',
    '- Rice Congee — easy on the system',
    '',
    'Section 2 — heading "AVOID TO PREVENT PUFFINESS":',
    '- Excess Salt — holds water in face and body',
    '- Late Night Eating — body can\'t process overnight',
    '- Alcohol — major cause of morning puffiness',
    '- Dairy — damp-forming, increases swelling',
    '- Cold Drinks — slow fluid processing',
]))

# ===== SERIES 5: EMOTION & SLEEP (3) =====

PROMPTS.append(("Series 5: Emotion & Sleep", 28, "Foods for Better Sleep", "Chinese Medicine Guide to Eat Your Way to Deep Sleep", [
    'Section 1 — heading "EAT TO SLEEP DEEPLY":',
    'Under subheading "Calm the Heart and Mind":',
    '- Lotus Seed — the classic calming sleep food',
    '- Lily Bulb — nourishes Heart Yin for deep rest',
    '- Red Dates — build Blood that anchors the mind',
    '- Longan Fruit — calms the spirit',
    '',
    'Under subheading "Warm and Soothing":',
    '- Warm Milk — gentle and comforting',
    '- Rice Congee (small bowl) — won\'t disturb sleep',
    '- Oatmeal — warm, complex carbs calm',
    '',
    'Under subheading "Nourish Yin":',
    '- Mulberry — nourishes Yin for nighttime',
    '- Wheat — traditional remedy for night restlessness',
    '- Goji Berries — gentle, nourishing',
    '',
    'Section 2 — heading "AVOID FOR BETTER SLEEP":',
    '- Coffee after noon — stays for 8 hours',
    '- Alcohol — disrupts deep sleep',
    '- Spicy Dinners — generate heat at night',
    '- Large Late Meals — body works instead of resting',
    '- Sugar Before Bed — crash causes 3 AM waking',
]))

PROMPTS.append(("Series 5: Emotion & Sleep", 29, "Foods for Stress Relief", "Chinese Medicine Foods to Unwind and Release Tension", [
    'Section 1 — heading "EAT TO RELEASE STRESS":',
    'Under subheading "Move Liver Energy":',
    '- Chrysanthemum Tea — cools frustration, calms head',
    '- Rose Petal Tea — moves Liver Qi, lifts mood',
    '- Mint Tea — light, refreshing, releases tension',
    '- Celery — moves energy downward',
    '',
    'Under subheading "Calms the Mind":',
    '- Lotus Seed — settles the spirit',
    '- Lily Bulb — cools and calms',
    '- Red Dates — grounds the mind through Blood',
    '- Wheat — traditional for emotional stability',
    '',
    'Under subheading "Nourish and Ground":',
    '- Bone Broth — deeply stabilizing',
    '- Warm Rice Congee — grounding comfort food',
    '- Walnuts — nourish the brain under stress',
    '',
    'Section 2 — heading "AVOID WHEN STRESSED":',
    '- Coffee — amplifies stress response',
    '- Alcohol — temporary mask, worse after',
    '- Sugar — crash worsens mood',
    '- Spicy Food — adds heat to frustration',
    '- Eating While Angry — impairs digestion',
]))

PROMPTS.append(("Series 5: Emotion & Sleep", 30, "Foods for Mood Swings", "Chinese Medicine Foods to Stabilize Your Emotions", [
    'Section 1 — heading "EAT FOR EMOTIONAL BALANCE":',
    'Under subheading "Soothe the Liver":',
    '- Rose Petal Tea — classic remedy for mood swings',
    '- Mint Tea — moves stuck energy, lifts mood',
    '- Chrysanthemum Tea — cools anger and frustration',
    '- Lemon Water — moves Liver Qi gently',
    '',
    'Under subheading "Nourish the Heart":',
    '- Lotus Seed — calms the spirit',
    '- Lily Bulb — stabilizes emotions',
    '- Red Dates — builds Blood for mental stability',
    '- Longan Fruit — warms and calms',
    '',
    'Under subheading "Stabilize Blood Sugar":',
    '- Sweet Potato — prevents sugar crashes',
    '- Millet Porridge — steady energy for mood',
    '- Chinese Yam — stabilizes and strengthens',
    '- Nuts (small amount) — protein for stability',
    '',
    'Section 2 — heading "AVOID FOR STABLE MOOD":',
    '- Excessive Sugar — crash creates irritability',
    '- Too Much Coffee — anxiety and mood spikes',
    '- Alcohol — rebound mood drops',
    '- Skipping Meals — low blood sugar = bad mood',
    '- Very Spicy Foods — add heat to emotions',
]))

# ===== SERIES 6: DIGESTION (3) =====

PROMPTS.append(("Series 6: Digestion", 31, "Chinese Medicine Digestion Guide", "How to Eat for a Happy Gut According to TCM", [
    'Section 1 — heading "EAT FOR HAPPY DIGESTION":',
    'Under subheading "Strengthen the Spleen":',
    '- Rice Congee — the most gut-friendly food',
    '- Millet Porridge — easiest grain to digest',
    '- Chinese Yam — strengthens Spleen directly',
    '- Sweet Potato — nourishes without burden',
    '',
    'Under subheading "Warm and Aid Digestion":',
    '- Ginger Tea — warms the digestive fire',
    '- Fennel Seeds — reduce gas naturally',
    '- Cardamom — eases fullness',
    '- Pumpkin — supports Spleen gently',
    '',
    'Under subheading "Easy Proteins":',
    '- Chicken Broth — deeply nourishing, easy',
    '- Steamed Fish — light, gentle protein',
    '- Tofu — easy plant protein',
    '',
    'Section 2 — heading "AVOID FOR GUT HEALTH":',
    '- Ice Cold Drinks — extinguish digestive fire',
    '- Raw Vegetables — hard on the Spleen',
    '- Dairy Products — damp-forming',
    '- Deep-Fried Foods — heaviest burden',
    '- Eating While Stressed — impairs all digestion',
]))

PROMPTS.append(("Series 6: Digestion", 32, "Foods That Reduce Gas and Bloating", "Chinese Medicine Natural Remedies for a Flat Stomach", [
    'Section 1 — heading "EAT FOR A FLAT STOMACH":',
    'Under subheading "Reduce Gas":',
    '- Fennel Seeds — the #1 anti-gas food',
    '- Ginger Tea — warms and moves digestion',
    '- Cardamom — reduces fullness and gas',
    '- White Radish — moves food downward',
    '',
    'Under subheading "Strengthen Digestion":',
    '- Chinese Yam — supports Spleen',
    '- Millet Porridge — easy, no gas',
    '- Cooked Carrots — gentle and warming',
    '- Rice Congee — no burden on gut',
    '',
    'Under subheading "Move Stuck Energy":',
    '- Mint Tea — moves Liver, reduces tension bloating',
    '- Papaya — enzymes break down food',
    '- Celery — light, moves energy',
    '',
    'Section 2 — heading "AVOID — THESE CAUSE GAS":',
    '- Raw Vegetables — ferment in weak digestion',
    '- Ice Cold Drinks — slow everything down',
    '- Carbonated Drinks — add air to gut',
    '- Excessive Beans (unsoaked) — gas-forming',
    '- Dairy — creates damp gas in many',
]))

PROMPTS.append(("Series 6: Digestion", 33, "Morning Digestion Routine", "Chinese Medicine Habits to Start Your Digestion Right", [
    'Section 1 — heading "MORNING DIGESTION GUIDE":',
    'Under subheading "First Thing (7-9 AM is Spleen Time)":',
    '- Warm Water with Lemon — gentle wake-up',
    '- Ginger Tea — starts the digestive fire',
    '- Rice Congee — warm, easy breakfast',
    '',
    'Under subheading "Build Morning Energy":',
    '- Millet Porridge — the best breakfast grain',
    '- Sweet Potato — steady energy start',
    '- Chinese Dates — natural Qi boost',
    '- Chinese Yam — strengthens for the day',
    '',
    'Under subheading "Support the Gut":',
    '- Chicken Broth — warm and restorative',
    '- Cooked Apple — soft, easy, non-acidic',
    '- Goji Berries in warm water — gentle nourishment',
    '',
    'Section 2 — heading "AVOID IN THE MORNING":',
    '- Ice Cold Water — shocks the Spleen',
    '- Coffee on Empty Stomach — burns borrowed energy',
    '- Skipping Breakfast — wastes peak Spleen time',
    '- Raw Cold Smoothies — too hard for morning',
    '- Sugary Pastries — crash by 10 AM',
]))

# ===== SERIES 7: TCM BASICS (4) =====

PROMPTS.append(("Series 7: TCM Basics", 34, "Cooling Foods vs Warming Foods", "Chinese Medicine Food Temperature Chart", [
    'Section 1 — heading "UNDERSTAND FOOD TEMPERATURES":',
    'Under subheading "WARMING FOODS (eat when cold, tired, or in winter)":',
    '- Lamb — the warmest meat',
    '- Ginger — the most common warming spice',
    '- Cinnamon — warm and sweet',
    '- Leeks and Chives — warm vegetables',
    '- Walnuts — warm and nourishing',
    '- Chestnuts — warm the Kidney',
    '',
    'Under subheading "COOLING FOODS (eat when hot, inflamed, or in summer)":',
    '- Mung Beans — the strongest cooling food',
    '- Watermelon — summer coolant',
    '- Cucumber — cooling and hydrating',
    '- Pear — cooling and moistening',
    '- Bitter Gourd — strong heat clearer',
    '- Lotus Root — cooling and crisp',
    '',
    'Under subheading "NEUTRAL FOODS (safe for everyone)":',
    '- Rice — neutral, gentle staple',
    '- Chinese Yam — neutral, strengthening',
    '- Carrot — neutral, nourishing',
    '',
    'Section 2 — heading "HOW TO CHOOSE":',
    '- Always cold? Choose more warming foods',
    '- Always hot? Choose more cooling foods',
    '- In summer? Add more cooling foods',
    '- In winter? Add more warming foods',
    '- Sick or weak? Stick to neutral and warm',
]))

PROMPTS.append(("Series 7: TCM Basics", 35, "What Is Dampness in Your Body?", "Chinese Medicine Guide to Understanding Dampness", [
    'Section 1 — heading "WHAT IS DAMPNESS?":',
    'Paragraph text: "In Chinese medicine, dampness is fluid your body cannot process. It feels heavy, sticky, and sluggish. A weak Spleen produces it, and poor habits worsen it."',
    '',
    'Section 2 — heading "SIGNS YOU MAY HAVE DAMPNESS":',
    '- Heavy feeling in body and limbs',
    '- Brain fog and clouded thinking',
    '- Sticky or sluggish bowel movements',
    '- Feeling of fullness and bloating',
    '- Weight that is hard to lose',
    '- Thick greasy coating on tongue',
    '- Cloudy or concentrated urine',
    '',
    'Section 3 — heading "FOODS THAT DRAIN DAMPNESS":',
    '- Adzuki Beans — the strongest',
    '- Coix Seed (Job\'s Tears)',
    '- Winter Melon',
    '- Lotus Leaf Tea',
    '- Barley Tea',
    '- White Radish',
    '- Celery',
    '',
    'Section 4 — heading "FOODS THAT CREATE DAMPNESS":',
    '- Dairy Products — the worst offender',
    '- Refined Sugar and Sweets',
    '- Deep-Fried Foods',
    '- Cold Drinks and Ice Water',
    '- Excessive Raw Foods',
    '',
    'Small footer at very bottom: "Take the free body type quiz -> myeasterntype.com"',
    '',
    STYLE_BLOCK,
]))

PROMPTS.append(("Series 7: TCM Basics", 36, "Foods That Drain Dampness", "Chinese Medicine Top 10 Damp-Removing Foods", [
    'Section 1 — heading "TOP 10 DAMP-DRAINING FOODS":',
    '',
    '- Adzuki Beans — the #1 damp-draining food',
    '- Coix Seed (Job\'s Tears) — drains and strengthens',
    '- Winter Melon — promotes fluid metabolism',
    '- Lotus Leaf Tea — reduces water weight',
    '- Barley Tea — gentle daily draining',
    '- White Radish — moves energy downward',
    '- Celery — natural diuretic',
    '- Corn Silk Tea — traditional edema remedy',
    '- Chen Pi (Aged Tangerine Peel) — transforms phlegm',
    '- Chinese Yam — strengthens Spleen to prevent damp',
    '',
    'Section 2 — heading "TOP DAMP-FORMING FOODS TO LIMIT":',
    '- Dairy — cheese, milk, ice cream',
    '- Refined Sugar — feeds dampness',
    '- Deep-Fried Foods — generate phlegm',
    '- Cold Drinks — slow fluid processing',
    '- Excessive Raw Foods — hard on Spleen',
    '',
    'Section 3 — heading "DAILY HABITS THAT HELP":',
    '- Exercise daily to sweat (dampness exits through sweat)',
    '- Eat all meals warm and cooked',
    '- Avoid drinking ice water',
    '- Keep your living space dry and ventilated',
    '',
    'Small footer at very bottom: "Take the free body type quiz -> myeasterntype.com"',
    '',
    STYLE_BLOCK,
]))

PROMPTS.append(("Series 7: TCM Basics", 37, "Yin and Yang Foods Chart", "Chinese Medicine Guide to Balancing Yin and Yang Through Diet", [
    'Section 1 — heading "YIN AND YANG FOOD CHART":',
    'Paragraph: "In Chinese medicine, Yin is cooling, moistening, and calming. Yang is warming, activating, and energizing. Balance is the goal."',
    '',
    'Under subheading "YIN FOODS (Cooling and Moistening)":',
    '- Watermelon — very cooling',
    '- Pear — moistens and cools',
    '- Mung Beans — strong cooling',
    '- Cucumber — hydrating and cool',
    '- Lotus Root — cooling and nourishing',
    '- Black Sesame — moistens and nourishes',
    '- Tofu — cooling plant protein',
    '- Mulberry — cool and nourishing',
    '',
    'Under subheading "YANG FOODS (Warming and Activating)":',
    '- Lamb — very warming',
    '- Ginger — warms the core',
    '- Cinnamon — warms and circulates',
    '- Leeks — warming vegetable',
    '- Walnuts — warm and nourishing',
    '- Chestnuts — warm the Kidney',
    '- Shrimp — warm and strengthening',
    '- Fennel — warming, aids digestion',
    '',
    'Under subheading "NEUTRAL FOODS (Balanced)":',
    '- Rice — gentle staple',
    '- Chinese Yam — neutral tonic',
    '- Carrot — neutral and sweet',
    '',
    'Section 2 — heading "HOW TO BALANCE":',
    '- Always hot and dry? Add more Yin foods',
    '- Always cold and tired? Add more Yang foods',
    '- Summer? More Yin foods to stay cool',
    '- Winter? More Yang foods to stay warm',
    '',
    'Small footer at very bottom: "Take the free body type quiz -> myeasterntype.com"',
    '',
    STYLE_BLOCK,
]))

# ===== SERIES 8: WOMEN'S HEALTH (3) =====

PROMPTS.append(("Series 8: Women's Health", 38, "Foods for Menopause Relief", "Chinese Medicine Cooling Foods for Hot Flashes and Night Sweats", [
    'Section 1 — heading "EAT FOR MENOPAUSE RELIEF":',
    'Under subheading "Cool Hot Flashes":',
    '- Mung Beans — clear empty heat',
    '- Tofu and Soy — cooling, phytoestrogen support',
    '- Cucumber — cooling and hydrating',
    '- Watermelon — quick cooling effect',
    '',
    'Under subheading "Nourish Kidney Yin":',
    '- Black Sesame Seeds — rebuild cooling reserve',
    '- Goji Berries — support Liver and Kidney',
    '- Mulberry — nourish Blood and Yin',
    '- Chinese Yam — gentle Kidney support',
    '',
    'Under subheading "Calm the Mind":',
    '- Lotus Seed — calms Heart, reduces anxiety',
    '- Lily Bulb — nourishes Heart Yin, aids sleep',
    '- Chrysanthemum Tea — cools frustration',
    '',
    'Section 2 — heading "AVOID DURING MENOPAUSE":',
    '- Spicy Foods — add fuel to hot flashes',
    '- Alcohol — generates internal heat',
    '- Excess Coffee — stimulates Heart and heat',
    '- Cinnamon and Lamb — too warming',
    '- Saunas — external heat worsens flashes',
]))

PROMPTS.append(("Series 8: Women's Health", 39, "Foods for PCOS Support", "Chinese Medicine Foods to Balance Hormones Naturally", [
    'Section 1 — heading "EAT FOR PCOS SUPPORT":',
    'Under subheading "Drain Dampness":',
    '- Adzuki Beans — drain damp, reduce cysts tendency',
    '- Coix Seed (Job\'s Tears) — drain and strengthen',
    '- Winter Melon — promote fluid metabolism',
    '- Lotus Leaf Tea — support weight management',
    '',
    'Under subheading "Support Kidney for Hormones":',
    '- Black Beans — nourish Kidney foundation',
    '- Chinese Yam — support Kidney and Spleen',
    '- Walnuts — nourish Kidney essence',
    '- Goji Berries — support Liver and Kidney',
    '',
    'Under subheading "Clear Heat for Skin":',
    '- Mung Beans — clear heat, help acne',
    '- Green Tea — moves Qi, antioxidant',
    '- Celery — moves energy, drains damp',
    '',
    'Section 2 — heading "AVOID WITH PCOS":',
    '- Dairy Products — most damp-forming',
    '- Refined Sugar — feeds damp-heat',
    '- Cold Drinks — weaken Spleen',
    '- Deep-Fried Foods — generate phlegm',
    '- Excessive Raw Foods — hard for Spleen',
]))

PROMPTS.append(("Series 8: Women's Health", 40, "Foods for PMS Mood Swings", "Chinese Medicine Foods to Calm Irritability Before Your Period", [
    'Section 1 — heading "EAT TO CALM PMS MOOD":',
    'Under subheading "Move Liver Energy":',
    '- Rose Petal Tea — the classic PMS mood remedy',
    '- Mint Tea — moves Qi, lifts mood',
    '- Chrysanthemum Tea — cools frustration',
    '- Lemon Water — moves Liver gently',
    '',
    'Under subheading "Calm the Heart":',
    '- Lotus Seed — settle the mind',
    '- Lily Bulb — cool and calm',
    '- Red Dates — ground through Blood',
    '',
    'Under subheading "Stabilize":',
    '- Banana — potassium for mood',
    '- Dark Chocolate (70%) — magnesium, small amount',
    '- Sweet Potato — prevents sugar crash mood',
    '- Millet Porridge — steady energy for emotions',
    '',
    'Section 2 — heading "AVOID BEFORE YOUR PERIOD":',
    '- Excess Coffee — amplifies irritability',
    '- Alcohol — worsens Liver stagnation',
    '- Too Much Sugar — crash worsens mood',
    '- Salty Foods — worsens bloating and tension',
    '- Ice Cold Drinks — worsens cramps and stagnation',
]))

# ===== SERIES 9: BODY TYPE FOOD MAPS (9) =====

PROMPTS.append(("Series 9: Body Type Food Maps", 41, "Foods for Qi Deficient Type", "Chinese Medicine Diet Guide for Low Energy Types", [
    'Section 1 — heading "IF YOU ARE ALWAYS TIRED, THIS IS FOR YOU":',
    'Paragraph: "Qi Deficient types feel exhausted after eating, have a soft voice, and catch colds easily. About 15% of people share this type."',
    '',
    'Under subheading "EAT MORE":',
    '- Sweet Potato — builds steady energy',
    '- Rice Congee — easiest energy source',
    '- Chinese Dates — classic Qi tonic',
    '- Ginger Tea — warms digestion',
    '- Chinese Yam — strengthens Spleen',
    '- Chicken Broth — deeply nourishing',
    '- Pumpkin — supports Spleen',
    '- Millet — most digestible grain',
    '',
    'Under subheading "LIMIT":',
    '- Ice Water — kills digestive fire',
    '- Raw Salads — too hard to process',
    '- Excessive Dairy — creates dampness',
    '- Deep-Fried Foods — drains energy',
    '- Too Much Coffee — borrows tomorrow\'s energy',
    '',
    'Small footer at very bottom: "Take the free body type quiz -> myeasterntype.com"',
    '',
    STYLE_BLOCK,
]))

PROMPTS.append(("Series 9: Body Type Food Maps", 42, "Foods for Yang Deficient Type", "Chinese Medicine Diet Guide for Always-Cold Types", [
    'Section 1 — heading "IF YOU ARE ALWAYS COLD, THIS IS FOR YOU":',
    'Paragraph: "Yang Deficient types feel cold especially in the lower body, have pale complexion, and frequent pale urination at night. About 10% of people."',
    '',
    'Under subheading "EAT MORE":',
    '- Lamb — the strongest warming meat',
    '- Walnuts — warm and nourish Kidney',
    '- Cinnamon — warms from the core',
    '- Ginger (dried) — deep internal warmth',
    '- Leeks and Chives — warming vegetables',
    '- Shrimp — warm and strengthening',
    '- Chestnuts — warm the Kidney',
    '- Fennel — warming, aids digestion',
    '',
    'Under subheading "LIMIT":',
    '- Ice Cold Drinks — the worst for Yang',
    '- Watermelon — too cooling',
    '- Mung Beans — very cooling',
    '- Excessive Raw Foods — require heat to process',
    '- Cucumber — cooling',
    '',
    'Small footer at very bottom: "Take the free body type quiz -> myeasterntype.com"',
    '',
    STYLE_BLOCK,
]))

PROMPTS.append(("Series 9: Body Type Food Maps", 43, "Foods for Yin Deficient Type", "Chinese Medicine Diet Guide for Hot and Dry Types", [
    'Section 1 — heading "IF YOU RUN HOT AND DRY, THIS IS FOR YOU":',
    'Paragraph: "Yin Deficient types have night sweats, dry mouth, hot flashes, and feel warm in the afternoon. About 10% of people."',
    '',
    'Under subheading "EAT MORE":',
    '- Black Sesame Seeds — moisten and cool',
    '- Goji Berries — nourish Liver and Kidney Yin',
    '- Pear — cooling and moistening',
    '- Mung Beans — clear internal heat',
    '- Mulberry — cool and nourishing',
    '- Lotus Seed — calm and moisten',
    '- Tofu — cooling plant protein',
    '- Duck — the one cooling meat',
    '',
    'Under subheading "LIMIT":',
    '- Spicy Foods — add fuel to empty heat',
    '- Coffee — heating and drying',
    '- Alcohol — generates heat rapidly',
    '- Lamb — too warming',
    '- Excessive Ginger — warming',
    '',
    'Small footer at very bottom: "Take the free body type quiz -> myeasterntype.com"',
    '',
    STYLE_BLOCK,
]))

PROMPTS.append(("Series 9: Body Type Food Maps", 44, "Foods for Phlegm Damp Type", "Chinese Medicine Diet Guide for Heavy and Sluggish Types", [
    'Section 1 — heading "IF YOU FEEL HEAVY AND HOLD WEIGHT, THIS IS FOR YOU":',
    'Paragraph: "Phlegm Damp types feel heavy, carry extra weight, and have sticky digestion. About 9% of people."',
    '',
    'Under subheading "EAT MORE":',
    '- Adzuki Beans — drain dampness strongly',
    '- Coix Seed (Job\'s Tears) — drain and strengthen',
    '- Winter Melon — promote fluid loss',
    '- White Radish — move and drain',
    '- Lotus Leaf Tea — reduce damp weight',
    '- Barley Tea — gentle daily draining',
    '- Celery — light and draining',
    '- Chinese Yam — strengthen Spleen',
    '',
    'Under subheading "LIMIT":',
    '- Dairy Products — the most damp-forming',
    '- Refined Sugar — feeds dampness',
    '- Cold Drinks — slow fluid processing',
    '- Deep-Fried Foods — generate phlegm',
    '- Excessive Raw Foods — burden the Spleen',
    '',
    'Small footer at very bottom: "Take the free body type quiz -> myeasterntype.com"',
    '',
    STYLE_BLOCK,
]))

PROMPTS.append(("Series 9: Body Type Food Maps", 45, "Foods for Damp Heat Type", "Chinese Medicine Diet Guide for Hot and Sticky Types", [
    'Section 1 — heading "IF YOU FEEL HOT AND STICKY WITH ACNE, THIS IS FOR YOU":',
    'Paragraph: "Damp Heat types run warm, feel sticky in humidity, and are prone to acne and skin issues. About 8% of people."',
    '',
    'Under subheading "EAT MORE":',
    '- Mung Beans — clear heat and damp together',
    '- Green Tea — clears heat, moves damp',
    '- Cucumber — cooling and draining',
    '- Bitter Gourd — clears heat strongly',
    '- Winter Melon — drains and cools',
    '- Lotus Root — cools and moves',
    '- Celery — drains, light',
    '- Pear — cools and moistens',
    '',
    'Under subheading "LIMIT":',
    '- Spicy Foods — add heat to the system',
    '- Alcohol — generates damp-heat directly',
    '- Deep-Fried Foods — create damp-heat',
    '- Dairy — damp-forming',
    '- Excessive Sweets — feed damp-heat',
    '',
    'Small footer at very bottom: "Take the free body type quiz -> myeasterntype.com"',
    '',
    STYLE_BLOCK,
]))

PROMPTS.append(("Series 9: Body Type Food Maps", 46, "Foods for Blood Stasis Type", "Chinese Medicine Diet Guide for Poor Circulation Types", [
    'Section 1 — heading "IF YOU BRUISE EASILY AND HAVE DULL PAIN, THIS IS FOR YOU":',
    'Paragraph: "Blood Stasis types have dark spots, varicose veins, dull chronic pain, and rough dry skin. About 7% of people."',
    '',
    'Under subheading "EAT MORE":',
    '- Black Fungus (Wood Ear) — improves Blood circulation',
    '- Hawthorn — moves Blood, digests fats',
    '- Saffron (small amount) — moves Blood beautifully',
    '- Rose Petal Tea — moves Blood and Qi',
    '- Peach Kernel — classic Blood mover',
    '- Chrysanthemum Tea — cools and circulates',
    '- Dark Cherries — support Blood health',
    '- Black Sesame Seeds — nourish Blood',
    '',
    'Under subheading "LIMIT":',
    '- Very Cold Foods — cold congeals Blood',
    '- Excessive Dairy — creates dampness in vessels',
    '- Deep-Fried Foods — sticky and stagnating',
    '- Ice Water — slows Blood flow',
    '- Refined Sugar — thickens Blood',
    '',
    'Small footer at very bottom: "Take the free body type quiz -> myeasterntype.com"',
    '',
    STYLE_BLOCK,
]))

PROMPTS.append(("Series 9: Body Type Food Maps", 47, "Foods for Qi Stagnant Type", "Chinese Medicine Diet Guide for Stressed and Tense Types", [
    'Section 1 — heading "IF YOU ARE TENSE, IRRITABLE, AND SIGH A LOT, THIS IS FOR YOU":',
    'Paragraph: "Qi Stagnant types hold stress in the body, feel frustrated easily, and have mood-related bloating. About 6% of people."',
    '',
    'Under subheading "EAT MORE":',
    '- Rose Petal Tea — the #1 remedy for this type',
    '- Mint Tea — moves stuck Liver energy',
    '- Chrysanthemum Tea — cools frustration',
    '- Celery — moves energy downward',
    '- Green Tea — light and moving',
    '- Lemon Water — gentle Liver mover',
    '- Radish (Daikon) — moves Qi downward',
    '- Kelp and Seaweed — soften stagnation',
    '',
    'Under subheading "LIMIT":',
    '- Heavy Greasy Foods — slow energy flow',
    '- Excessive Alcohol — worsens Liver stagnation',
    '- Too Much Coffee — overstimulates',
    '- Very Spicy Foods — convert stagnation to fire',
    '- Fried Foods — create sluggish energy',
    '',
    'Small footer at very bottom: "Take the free body type quiz -> myeasterntype.com"',
    '',
    STYLE_BLOCK,
]))

PROMPTS.append(("Series 9: Body Type Food Maps", 48, "Foods for Sensitive Type", "Chinese Medicine Diet Guide for Allergy-Prone Types", [
    'Section 1 — heading "IF YOU SNEEZE EASILY AND HAVE SENSITIVE SKIN, THIS IS FOR YOU":',
    'Paragraph: "Sensitive types react easily to pollen, dust, and certain foods. They may have allergic rhinitis or eczema. About 3% of people."',
    '',
    'Under subheading "EAT MORE":',
    '- Chinese Dates — build defensive energy',
    '- Chinese Yam — strengthen Lung and Spleen',
    '- Lotus Seed — support immune foundation',
    '- White Fungus — moisten Lung, reduce reactions',
    '- Lily Bulb — calm and moisten',
    '- Pear — protect and moisten Lung',
    '- Honey (local, raw) — traditional allergy support',
    '- Rice Congee — gentle, non-reactive',
    '',
    'Under subheading "LIMIT":',
    '- Dairy Products — increase phlegm and reactivity',
    '- Excessive Sugar — feeds inflammation',
    '- Cold Drinks — weaken Lung defense',
    '- Deep-Fried Foods — create phlegm',
    '- Known Allergen Foods — listen to your body',
    '',
    'Small footer at very bottom: "Take the free body type quiz -> myeasterntype.com"',
    '',
    STYLE_BLOCK,
]))

PROMPTS.append(("Series 9: Body Type Food Maps", 49, "Foods for Balanced Type", "Chinese Medicine Diet Guide for the Balanced Constitution", [
    'Section 1 — heading "IF YOU FEEL PRETTY GOOD OVERALL, THIS IS FOR YOU":',
    'Paragraph: "Balanced types sleep well, have good energy, digest comfortably, and feel emotionally stable. About 32% of people. The goal is to maintain this balance."',
    '',
    'Under subheading "EAT TO MAINTAIN BALANCE":',
    '- Seasonal Vegetables — eat with the seasons',
    '- Rice and Millet — balanced staples',
    '- Chinese Dates — gentle daily tonic',
    '- Goji Berries — daily maintenance',
    '- Chinese Yam — gentle strengthening',
    '- Moderate Variety — balance all food groups',
    '',
    'Under subheading "SEASONAL ADJUSTMENT":',
    '- Summer: add mung beans, cucumber, watermelon',
    '- Winter: add lamb, ginger, cinnamon',
    '- Spring: add mint, greens, chrysanthemum',
    '- Autumn: add pear, white fungus, honey',
    '',
    'Under subheading "DAILY HABITS":',
    '- Eat at regular times',
    '- Sleep before 11 PM',
    '- Exercise moderately',
    '- Adapt diet to seasons',
    '',
    'Section 2 — heading "EVEN BALANCED TYPES SHOULD LIMIT":',
    '- Excessive Ice Cold Drinks',
    '- Too Much Processed Food',
    '- Irregular Meal Times',
    '- Overeating',
    '',
    'Small footer at very bottom: "Take the free body type quiz -> myeasterntype.com"',
    '',
    STYLE_BLOCK,
]))

# ===== GENERATE DOCUMENT =====

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title page
title = doc.add_heading('Pinterest Pin Prompts — EastType', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('49 Information-Rich Food Guide Panels\nUnified Style Guide for Image Generation')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0xC9, 0xA3, 0x55)

doc.add_paragraph()

# Style guide section
doc.add_heading('Unified Style Guide (Applies to All 49 Pins)', level=1)

p = doc.add_paragraph()
p.add_run('Image Size: ').bold = True
p.add_run('1000x1500px (2:3 vertical ratio)')

p = doc.add_paragraph()
p.add_run('Color Palette: ').bold = True
p.add_run('Cream background (#FAF6F0), Gold accents (#C9A355), Warm brown text (#2A1F14)')

p = doc.add_paragraph()
p.add_run('Font: ').bold = True
p.add_run('Playfair Display (serif) for titles, clean sans-serif for body')

p = doc.add_paragraph()
p.add_run('Style Keywords: ').bold = True
p.add_run('Warm wellness aesthetic, subtle Chinese medicine heritage feel, small flat food illustrations beside key items, soft shadows, generous whitespace, organized in clear sections')

p = doc.add_paragraph()
p.add_run('Footer: ').bold = True
p.add_run('Very small, unobtrusive text at bottom edge: "Take the free body type quiz -> myeasterntype.com"')

p = doc.add_paragraph()
p.add_run('Language: ').bold = True
p.add_run('All text in English only. No Chinese characters on images.')

p = doc.add_paragraph()
p.add_run('Vibe: ').bold = True
p.add_run('Professional, saveable, information-rich. Users should want to save this as a reference card.')

doc.add_page_break()

# Table of contents
doc.add_heading('Contents', level=1)
toc_items = [
    'Series 1: Menstrual Cycle Food Guides (4 pins)',
    'Series 2: Daily Symptoms Food Guides (15 pins)',
    'Series 3: Seasonal Health (4 pins)',
    'Series 4: Skin & Beauty (4 pins)',
    'Series 5: Emotion & Sleep (3 pins)',
    'Series 6: Digestion (3 pins)',
    'Series 7: TCM Basics (4 pins)',
    'Series 8: Women\'s Health (3 pins)',
    'Series 9: Body Type Food Maps (9 pins)',
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number')

doc.add_page_break()

# Generate each prompt
current_series = ""
for series, num, title_text, subtitle_text, sections in PROMPTS:
    # Series heading (only when series changes)
    if series != current_series:
        current_series = series
        doc.add_heading(series, level=1)

    # Pin number and title
    doc.add_heading(f'Pin #{num}: {title_text}', level=2)

    # Build the full prompt text
    if isinstance(sections, list) and len(sections) > 0 and isinstance(sections[0], str):
        # Standard format
        prompt_text = make_prompt(title_text, subtitle_text, ['\n'.join(sections)])
    else:
        prompt_text = make_prompt(title_text, subtitle_text, sections)

    # Some entries already have full prompt text built into sections
    # Check if sections contains the full STYLE_BLOCK already
    full_text = '\n'.join(sections) if isinstance(sections, list) else str(sections)

    # Check if this is a pre-built prompt (has STYLE_BLOCK in it)
    has_style = any('STYLE_BLOCK' in str(s) for s in sections) if isinstance(sections, list) else 'STYLE_BLOCK' in str(sections)

    if has_style:
        # Replace STYLE_BLOCK with actual style text
        full_text = full_text.replace('STYLE_BLOCK', STYLE_BLOCK)
        prompt_para = doc.add_paragraph()
        run = prompt_para.add_run(full_text)
        run.font.name = 'Consolas'
        run.font.size = Pt(10)
    else:
        # Build the prompt from components
        prompt_lines = []
        prompt_lines.append(f'Create a vertical Pinterest infographic (1000x1500px, 2:3 ratio)')
        prompt_lines.append(f'about {title_text.lower()}.')
        prompt_lines.append('')
        prompt_lines.append(f'Title in elegant serif font: "{title_text}"')
        if subtitle_text:
            prompt_lines.append(f'Subtitle in smaller text: "{subtitle_text}"')
        prompt_lines.append('')

        for sec in sections:
            prompt_lines.append(sec)

        prompt_lines.append('')
        prompt_lines.append(FOOTER)
        prompt_lines.append('')
        prompt_lines.append(STYLE_BLOCK)

        full_prompt = '\n'.join(prompt_lines)

        prompt_para = doc.add_paragraph()
        run = prompt_para.add_run(full_prompt)
        run.font.name = 'Consolas'
        run.font.size = Pt(10)

    doc.add_paragraph('---')
    doc.add_paragraph()

# Save
output_path = os.path.join(os.getcwd(), 'Pinterest_Pin_Prompts_49.docx')
doc.save(output_path)
print(f'Document saved: {output_path}')
print(f'Total prompts: {len(PROMPTS)}')
