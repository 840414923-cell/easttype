from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

title = doc.add_heading("Homepage Image Prompts — 4 Content Blocks", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("For image2 image generation. Square 600x600px. Match website style: warm cream, gold accents, clean wellness aesthetic.")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0xC9, 0xA3, 0x55)

doc.add_paragraph()

doc.add_heading("General Visual Style (applies to all 4 images)", level=1)

style_items = [
    ("Resolution", "600x600px (1:1 square ratio)"),
    ("Background", "Warm cream (#F5EFE6) or very light warm beige. Soft, clean, minimal."),
    ("Color Palette", "Cream (#F5EFE6), gold accent (#C9A355), warm brown (#2A1F14), soft sage green (#8E9B90), muted terracotta (#C75B3A). Natural, earthy, warm."),
    ("Overall Feel", "Premium Chinese medicine wellness brand. Modern, clean, approachable. NOT clinical. NOT old-fashioned. Think: high-end tea brand meets wellness lifestyle magazine."),
    ("Consistency", "All 4 images must share the same visual language: same background tone, same illustration style, same color palette. They should feel like a cohesive set."),
]

for label, value in style_items:
    p = doc.add_paragraph()
    p.add_run(f"{label}: ").bold = True
    p.add_run(value)

doc.add_paragraph()
doc.add_page_break()

# Image 1: Body Types
doc.add_heading("Image 1: Body Types", level=1)
doc.add_paragraph("Position: right side of first content block (left text, right image)", style="Intense Quote")
doc.add_paragraph()

prompt1 = """Create a square image (600x600px, 1:1 ratio) for a Chinese medicine wellness website.

Theme: 9 Body Types concept illustration.

Background: warm cream color (#F5EFE6), clean and minimal. A very subtle golden radial glow from the center at 8% opacity.

Center composition: 9 small circular icons arranged in a 3x3 grid pattern in the center of the image. Each circle is about 60px diameter, with a thin gold border (#C9A355, 2px). The circles are spaced evenly with 20px gaps.

Each circle contains a simple, elegant flat-illustration icon representing one of the 9 Chinese medicine body types:
1. Top-left: a calm still lake ( Balanced type)
2. Top-center: a gentle breeze / curved wind lines (Qi Deficient)
3. Top-right: a warm sun with rays (Yang Deficient)
4. Middle-left: a crescent moon with small flame (Yin Deficient)
5. Middle-center: a rounded stone with moss dots (Phlegm Damp)
6. Middle-right: a stylized sun with rain drops (Damp Heat)
7. Bottom-left: a curved stream with a small ice crystal (Blood Stasis)
8. Bottom-center: a cloud with swirling lines (Qi Stagnant)
9. Bottom-right: a small orchid flower (Sensitive type)

All icons are drawn in thin gold lines (#C9A355) on the cream background, flat illustration style, consistent line weight (1.5px), no fills, no shading. Clean and minimal like a luxury brand icon set.

Below the 3x3 grid, centered, small text in Playfair Display serif font, gold color (#C9A355): "9 Body Types"

Style: premium wellness brand, minimal, elegant, modern. The overall feeling should be like an expensive wellness app icon set or a luxury Chinese medicine brand logo. No photographic elements. Flat line illustration only. Generous whitespace around the grid."""

doc.add_paragraph(prompt1)

doc.add_paragraph()
doc.add_page_break()

# Image 2: Symptoms
doc.add_heading("Image 2: Symptoms", level=1)
doc.add_paragraph("Position: left side of second content block (left image, right text)", style="Intense Quote")
doc.add_paragraph()

prompt2 = """Create a square image (600x600px, 1:1 ratio) for a Chinese medicine wellness website.

Theme: Chinese medicine symptom diagnosis illustration.

Background: warm cream color (#F5EFE6), clean and minimal.

Center composition: a large circular frame (300px diameter) in the center, drawn with a thin gold border (#C9A355, 2px). Inside the circle, a flat-illustration of a tongue diagnosis concept: a simplified, elegant side-profile of a face in thin gold lines, with a small tongue visible. The tongue has subtle color zones indicated by very light watercolor-style washes: pale pink center, thin white coating area, slightly red tip. The face is drawn in a minimal, modern way, NOT a medical diagram. Think: artistic, editorial illustration.

Around the large circle, 6 small gold dots (#C9A355, 8px each) are placed at even intervals on the circle's edge, like a clock face. Each dot has a tiny thin gold line connecting it to a small label outside the circle. The labels are tiny text in brown (#2A1F14, 9pt):

Top: "Sleep"
Top-right: "Energy"
Bottom-right: "Digestion"
Bottom: "Mood"
Bottom-left: "Skin"
Top-left: "Temperature"

Below the circle, centered, small text in Playfair Display serif font, gold (#C9A355): "70 Symptom Guides"

Style: premium wellness brand, editorial illustration feel. The tongue diagnosis should feel artistic and intriguing, NOT clinical or medical. Think: Kinfolk magazine meets Chinese medicine. Thin gold lines, subtle watercolor touches, generous whitespace."""

doc.add_paragraph(prompt2)

doc.add_paragraph()
doc.add_page_break()

# Image 3: Food Guides
doc.add_heading("Image 3: Food Guides", level=1)
doc.add_paragraph("Position: right side of third content block (left text, right image)", style="Intense Quote")
doc.add_paragraph()

prompt3 = """Create a square image (600x600px, 1:1 ratio) for a Chinese medicine wellness website.

Theme: Chinese medicine food therapy illustration.

Background: warm cream color (#F5EFE6), clean and minimal.

Center composition: a flat-illustration still-life arrangement of 6 key Chinese medicine foods, arranged in a loose circular composition in the center. Each food item is drawn in a modern, elegant flat illustration style with subtle warm-toned colors (not realistic photography, but stylized illustration):

1. Top-left: 3 red dates (jujube) - small oval shapes in warm red-brown (#9B6B5A)
2. Top-right: a small handful of goji berries - tiny oval shapes in bright orange-red (#C75B3A)
3. Middle-left: a ginger root - irregular shape in warm beige (#D4C9B8) with gold outline
4. Middle-right: a small bowl of rice congee - cream colored bowl with white porridge, steam rising in thin gold lines
5. Bottom-left: a sweet potato - elongated oval in warm orange (#E0C878)
6. Bottom-right: a chrysanthemum flower - simple 8-petal floral shape in pale gold (#C9A355)

All items are connected by a very thin gold circular line (#C9A355, 1px, 30% opacity) that flows behind them, tying the composition together.

Each item has a tiny gold label below it in 8pt Playfair Display font: "Red Dates", "Goji", "Ginger", "Congee", "Sweet Potato", "Chrysanthemum".

Below the arrangement, centered, small text in Playfair Display serif font, gold (#C9A355): "Food Therapy"

Style: premium wellness brand, modern flat illustration. Think: high-end food magazine illustration meets Chinese medicine. Warm, inviting, appetizing but elegant. NOT a medical herb chart. NOT photography. Generous whitespace."""

doc.add_paragraph(prompt3)

doc.add_paragraph()
doc.add_page_break()

# Image 4: Wellness Guides
doc.add_heading("Image 4: Wellness Guides", level=1)
doc.add_paragraph("Position: left side of fourth content block (left image, right text)", style="Intense Quote")
doc.add_paragraph()

prompt4 = """Create a square image (600x600px, 1:1 ratio) for a Chinese medicine wellness website.

Theme: Chinese medicine concepts and patterns illustration.

Background: warm cream color (#F5EFE6), clean and minimal.

Center composition: a yin-yang inspired circular arrangement in the center (280px diameter). But instead of the traditional black-and-white yin-yang symbol, use a modern interpretation:

The circle is divided into two halves by a gentle S-curve (like the yin-yang). The left half is filled with a very subtle warm gold gradient (#C9A355 to #E0C878, 15% opacity). The right half is filled with a very subtle warm brown gradient (#2A1F14 to #6B5D4D, 10% opacity). The dividing S-curve is drawn in solid gold (#C9A355, 2px).

Around the circle, 5 small concept icons are placed at equal intervals outside the circle, connected by thin gold lines (1px, 40% opacity) to the circle's edge:

1. Top: a small lotus flower (8 petals) in gold line
2. Top-right: a small Qi symbol (stylized energy swirl) in gold line
3. Bottom-right: a small meridian line (curved dotted line) in gold line
4. Bottom-left: a small 5-pointed star (representing Five Elements) in gold line
5. Top-left: a small leaf (representing food therapy) in gold line

Each icon is about 30px, drawn in thin gold lines (#C9A355, 1.5px), no fill.

Below the composition, centered, small text in Playfair Display serif font, gold (#C9A355): "43 Wellness Guides"

Style: premium wellness brand, modern and abstract. The yin-yang should feel contemporary, not traditional. Think: luxury wellness brand logo meets editorial design. Clean, elegant, meaningful. Generous whitespace."""

doc.add_paragraph(prompt4)

doc.add_page_break()

# Summary
doc.add_heading("Summary", level=1)

doc.add_paragraph()
doc.add_paragraph("All 4 images share: same cream background, same gold accent color, same flat illustration style, same Playfair Display font for labels. They should look like a cohesive set from the same brand.", style="Intense Quote")

table = doc.add_table(rows=1, cols=4)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
hdr[0].text = "Image"
hdr[1].text = "Theme"
hdr[2].text = "Position"
hdr[3].text = "Key Element"

data = [
    ("1", "Body Types", "Right side", "3x3 grid of 9 circular icons"),
    ("2", "Symptoms", "Left side", "Tongue diagnosis in circular frame"),
    ("3", "Food Guides", "Right side", "6 food illustrations in circle"),
    ("4", "Wellness Guides", "Left side", "Modern yin-yang with concept icons"),
]
for row_data in data:
    row = table.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

output = os.path.join(os.getcwd(), "Homepage_Image_Prompts_4.docx")
doc.save(output)
print(f"Saved: {output}")
