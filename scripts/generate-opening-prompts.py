from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

title = doc.add_heading("Opening Animation Image Prompts — 3 Versions", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("For image2 image generation. Text can be rendered on image.")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0xC9, 0xA3, 0x55)

doc.add_paragraph()

doc.add_heading("General Visual Style (applies to all 3 versions)", level=1)

style_items = [
    ("Resolution", "1000x1500px (2:3 vertical, for short video)"),
    ("Background", "Warm cream (#F5EFE6) with subtle golden radial glow from center. Faint Chinese medicine herb illustrations (ginger root, goji berries, chrysanthemum flowers) as watermark-style background decoration at 5% opacity."),
    ("Color Palette", "Cream background (#F5EFE6), gold accent (#C9A355), warm brown text (#2A1F14), soft red for highlights (#C75B3A). Clean, premium wellness brand aesthetic."),
    ("Host Character", "Young East Asian woman, mid-20s to early-30s, wearing a modern minimal olive-green or warm beige top. Natural makeup. Hair in a simple style. Friendly but confident expression. Professional wellness expert vibe, not clinical. Think: modern TCM practitioner on social media."),
    ("Text Style", "Playfair Display serif font for main text. Clean, large, centered. Gold color (#C9A355) for the hook text. White or cream rounded rectangle behind text for readability."),
    ("Overall Feel", "Premium Chinese medicine wellness brand. Warm, trustworthy, modern. NOT traditional/old-fashioned. NOT clinical/medical. Think Headspace meets a high-end tea brand."),
]

for label, value in style_items:
    p = doc.add_paragraph()
    p.add_run(f"{label}: ").bold = True
    p.add_run(value)

doc.add_paragraph()

doc.add_page_break()

# Version A
doc.add_heading("Version A: Contrarian Hook", level=1)

doc.add_paragraph("Hook Text: Most people get this completely wrong.", style="Intense Quote")

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Image Prompt for image2:").bold = True
p.add_run(" (copy everything below)")

doc.add_paragraph()

prompt_a = """Create a vertical image (1000x1500px, 2:3 ratio) for a short video opening frame.

Background: warm cream color (#F5EFE6) with a soft golden radial glow emanating from the center. Very faint illustrations of traditional Chinese medicine herbs (ginger root, goji berries, chrysanthemum flowers) as delicate watermark-style decorations at 5% opacity, barely visible.

Host: A young East Asian woman in her late 20s, wearing a modern minimal olive-green top. Natural makeup, simple hair style. She is looking directly at the camera with a confident, slightly challenging expression. One eyebrow slightly raised, as if saying "you probably don't know this." She is positioned on the right side of the frame, upper body visible.

Text overlay on the left side, in large elegant Playfair Display serif font, gold color (#C9A355):
"Most people get this completely wrong."

Below the text, a small decorative element: a thin gold line (1px) with a tiny diamond symbol (diamond) in the center.

At the very bottom, small text: "myeasterntype.com" in brown (#2A1F14), 10pt size.

Style: premium wellness brand, warm and modern, clean composition with generous whitespace. Professional but approachable. NOT clinical, NOT old-fashioned. Think: modern Chinese medicine meets high-end lifestyle brand."""

doc.add_paragraph(prompt_a)

doc.add_paragraph()
doc.add_paragraph("Chinese translation of hook: da duo shu ren wan quan xuan cuo le (most people get it completely wrong)")
doc.add_paragraph("Animation note: Host holds still for 1.5 seconds, then slight head tilt at the word 'wrong.' Total 2 seconds.")

doc.add_page_break()

# Version B
doc.add_heading("Version B: Curiosity Gap", level=1)

doc.add_paragraph("Hook Text: Chinese medicine says the opposite of what you think.", style="Intense Quote")

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Image Prompt for image2:").bold = True
p.add_run(" (copy everything below)")

doc.add_paragraph()

prompt_b = """Create a vertical image (1000x1500px, 2:3 ratio) for a short video opening frame.

Background: warm cream color (#F5EFE6) with a soft golden radial glow from center. Very faint illustrations of Chinese medicine herbs (lotus seed pod, red dates, yam root) as delicate watermark decorations at 5% opacity.

Host: A young East Asian woman in her late 20s, wearing a modern minimal warm beige top. Natural makeup, simple hair. She is leaning forward slightly toward the camera with a mysterious, knowing smile. Eyes slightly narrowed as if sharing a secret. Positioned on the right side, upper body visible.

Text overlay on the left side, in large elegant Playfair Display serif font, gold color (#C9A355):
"Chinese medicine says the opposite of what you think."

The word "opposite" is highlighted in a slightly larger size or warmer gold tone to emphasize the twist.

Below the text, a small decorative element: a thin gold line with a tiny lotus symbol in the center.

At the very bottom, small text: "myeasterntype.com" in brown (#2A1F14), 10pt size.

Style: premium wellness brand, warm and intriguing, clean composition. The host expression should convey "I know something you don't." Modern Chinese medicine meets lifestyle brand aesthetic."""

doc.add_paragraph(prompt_b)

doc.add_paragraph()
doc.add_paragraph("Animation note: Host leans in slowly, smile appears at the word 'opposite.' Total 2-3 seconds.")

doc.add_page_break()

# Version C
doc.add_heading("Version C: Challenge", level=1)

doc.add_paragraph("Hook Text: 9 out of 10 people choose the wrong answer. Are you the 1?", style="Intense Quote")

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Image Prompt for image2:").bold = True
p.add_run(" (copy everything below)")

doc.add_paragraph()

prompt_c = """Create a vertical image (1000x1500px, 2:3 ratio) for a short video opening frame.

Background: warm cream color (#F5EFE6) with a soft golden radial glow from center. Very faint illustrations of Chinese medicine herbs (walnut, black sesame, dried tangerine peel) as delicate watermark decorations at 5% opacity.

Host: A young East Asian woman in her late 20s, wearing a modern minimal deep teal or forest green top. Natural makeup, simple hair. She is holding up three fingers toward the camera in a confident, playful gesture. Expression is challenging and inviting, like a game show host who knows the answer. Direct eye contact. Positioned center-right, upper body visible.

Text overlay on the upper-left area, in large elegant Playfair Display serif font, gold color (#C9A355):
"9 out of 10 people choose the wrong answer."

Below that line, slightly smaller text:
"Are you the 1?"

The number "1" is highlighted in a warm red accent (#C75B3A) to make it pop.

Below the text, a small decorative element: a thin gold line with a tiny diamond symbol.

At the very bottom, small text: "myeasterntype.com" in brown (#2A1F14), 10pt size.

Style: premium wellness brand, warm and engaging, gamified feel. The host should look like she's inviting you to a challenge. Modern Chinese medicine aesthetic."""

doc.add_paragraph(prompt_c)

doc.add_paragraph()
doc.add_paragraph("Animation note: Host raises 3 fingers on '9 out of 10', points to camera on 'Are you the 1.' Total 2-3 seconds.")

doc.add_page_break()

# Ending section
doc.add_heading("Bonus: Ending Animation Image Prompts (2 versions)", level=1)

doc.add_paragraph("Same host, same background, different text overlay.", style="Intense Quote")

doc.add_paragraph()

doc.add_heading("Ending Version A: Standard", level=2)

doc.add_paragraph("Text: Answer in comments. First correct answer wins a free body type quiz code!")

doc.add_paragraph()

prompt_ea = """Create a vertical image (1000x1500px, 2:3 ratio) for a short video ending frame.

Background: same warm cream (#F5EFE6) with golden glow. Same faint herb illustrations.

Host: Same young East Asian woman, now with a warm welcoming smile, both hands gesturing toward the comment section (pointing downward or making a typing gesture). Expression: encouraging and inviting.

Text overlay, centered, in Playfair Display serif font:
Line 1 (gold, #C9A355, large): "Answer in comments."
Line 2 (brown, #2A1F14, medium): "First correct answer wins a free"
Line 3 (gold, #C9A355, medium, bold): "body type quiz code!"

At the very bottom: "myeasterntype.com" small text.

Style: same premium wellness aesthetic. Clean, warm, inviting."""

doc.add_paragraph(prompt_ea)

doc.add_paragraph()

doc.add_heading("Ending Version B: Urgent", level=2)

doc.add_paragraph("Text: Comment your answer NOW. Correct answers get a free quiz code in your DM!")

doc.add_paragraph()

prompt_eb = """Create a vertical image (1000x1500px, 2:3 ratio) for a short video ending frame.

Background: same warm cream (#F5EFE6) with golden glow. Same faint herb illustrations.

Host: Same young East Asian woman, now with an excited, urgent expression. One hand pointing directly at the camera, other hand holding a small gift box or envelope icon (representing the code). Expression: "hurry up and comment!"

Text overlay, centered, in Playfair Display serif font:
Line 1 (gold, #C9A355, large, bold): "Comment NOW!"
Line 2 (brown, #2A1F14, medium): "Correct answers get a free"
Line 3 (gold, #C9A355, medium): "quiz code in your DM!"

Add a small gift icon (gift) next to the text.

At the very bottom: "myeasterntype.com" small text.

Style: same premium wellness aesthetic but with more energy and urgency. Clean, warm, exciting."""

doc.add_paragraph(prompt_eb)

doc.add_page_break()

# Summary
doc.add_heading("Summary: What You Need to Create", level=1)

doc.add_paragraph()

summary = [
    ("1 opening image", "Choose A, B, or C. Make it once, use for all 20 videos."),
    ("1 ending image", "Choose A or B. Make it once, use for all 20 videos."),
    ("20 question images", "Each contains question text + 3 options. Use the content from Short_Video_Scripts_Final_20.docx."),
    ("Total images needed", "22 images (1 opening + 1 ending + 20 questions)"),
    ("Video assembly", "Opening animation (2s) + Question image (7s) + Ending animation (3s) = 12-15s per video"),
]

table = doc.add_table(rows=1, cols=2)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
hdr[0].text = "Item"
hdr[1].text = "Description"
for item, desc in summary:
    row = table.add_row().cells
    row[0].text = item
    row[1].text = desc

doc.add_paragraph()
doc.add_paragraph("Tip: Keep the host character, background, and color palette CONSISTENT across all images. This builds brand recognition.", style="Intense Quote")

output = os.path.join(os.getcwd(), "Opening_Ending_Animation_Prompts.docx")
doc.save(output)
print(f"Saved: {output}")
