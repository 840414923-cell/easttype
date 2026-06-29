from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json, os

with open(os.path.join(os.path.dirname(__file__), "quiz-data.json"), "r", encoding="utf-8") as f:
    questions = json.load(f)

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# Title
title = doc.add_heading("Short Video Script — 20 Episodes", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("English script for AI avatar + Chinese reference below each section")
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0xC9, 0xA3, 0x55)
doc.add_paragraph()

# Structure guide
doc.add_heading("Video Structure (15 seconds)", level=1)
guide = [
    ("0-2s", "Opening hook (avatar speaks to camera)"),
    ("2-3s", "Question appears (text + voice)"),
    ("3-9s", "3 options shown one by one (2s each)"),
    ("9-11s", "Options freeze + 'Comment your answer'"),
    ("11-13s", "'Free quiz code for correct answers'"),
    ("13-15s", "Logo + myeasterntype.com"),
]
for time, desc in guide:
    p = doc.add_paragraph()
    p.add_run(f"{time}: ").bold = True
    p.add_run(desc)

doc.add_page_break()

# Opening options
doc.add_heading("Opening Hook Options", level=1)
openings = [
    ("Option A (Recommended)", "Most people get this completely wrong."),
    ("Option B", "Chinese medicine says the opposite of what you think."),
    ("Option C", "9 out of 10 people choose the wrong answer. Are you the 1?"),
]
for name, text in openings:
    p = doc.add_paragraph()
    p.add_run(f"{name}: ").bold = True
    p.add_run(text)
    doc.add_paragraph()

# Ending options
doc.add_heading("Ending Options", level=1)
endings = [
    ("Option A (Recommended)", "Answer in comments. First correct answer wins a free body type quiz code!"),
    ("Option B", "Comment your answer NOW. Correct answers get a free quiz code in your DM!"),
]
for name, text in endings:
    p = doc.add_paragraph()
    p.add_run(f"{name}: ").bold = True
    p.add_run(text)
    doc.add_paragraph()

doc.add_page_break()

# Generate each episode
for q in questions:
    doc.add_heading(f"Episode {q['n']}", level=1)

    # Full script in English
    doc.add_heading("Full Script (English)", level=2)

    p = doc.add_paragraph()
    p.add_run("[0-2s Opening Hook]\n").bold = True
    p.add_run("Most people get this completely wrong.")

    p = doc.add_paragraph()
    p.add_run("[2-3s Question]\n").bold = True
    p.add_run(q["q_en"])

    p = doc.add_paragraph()
    p.add_run("[3-9s Options]\n").bold = True
    p.add_run(q["a"] + "\n")
    p.add_run(q["b"] + "\n")
    p.add_run(q["c"])

    p = doc.add_paragraph()
    p.add_run("[9-13s Call to Action]\n").bold = True
    p.add_run("Comment your answer below. First correct answer wins a free body type quiz code!")

    p = doc.add_paragraph()
    p.add_run("[13-15s Brand]\n").bold = True
    p.add_run("myeasterntype.com")

    # Chinese reference
    doc.add_heading("Chinese Reference", level=2)

    p = doc.add_paragraph()
    p.add_run("问题: ").bold = True
    p.add_run(q["q_zh"])

    p = doc.add_paragraph()
    p.add_run("选项:\n").bold = True
    p.add_run(q["a"] + "\n")
    p.add_run(q["b"] + "\n")
    p.add_run(q["c"])

    p = doc.add_paragraph()
    p.add_run("答案: ").bold = True
    ans_run = p.add_run(q["ans"])
    ans_run.bold = True
    ans_run.font.color.rgb = RGBColor(0xC9, 0xA3, 0x55)

    p = doc.add_paragraph()
    p.add_run("原因: ").bold = True
    p.add_run(q["why_zh"])

    # Comment section template
    doc.add_heading("Comment Section Reply", level=2)
    comment_en = f"Answer: {q['ans']} ✓ {q['why_en']} First correct answer wins a free body type quiz code! Check your DM 🎁"
    p = doc.add_paragraph()
    r = p.add_run(comment_en)
    r.font.size = Pt(10)

    doc.add_paragraph("─" * 40)
    doc.add_paragraph()

output = os.path.join(os.path.dirname(__file__), "..", "Short_Video_Scripts_20.docx")
doc.save(output)
print(f"Saved: {output}")
print(f"Episodes: {len(questions)}")
