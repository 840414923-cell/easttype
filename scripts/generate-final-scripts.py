from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json, os

with open(os.path.join("scripts", "quiz-data.json"), "r", encoding="utf-8") as f:
    questions = json.load(f)

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

title = doc.add_heading("Short Video Scripts — 20 Episodes", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Fixed Opening Animation + Question Image + Fixed Ending")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0xC9, 0xA3, 0x55)

doc.add_paragraph()

doc.add_heading("FIXED OPENING ANIMATION (Same for all videos)", level=1)

doc.add_paragraph("This is a pre-made animation used at the start of EVERY video. It does NOT mention any specific question.", style="Intense Quote")

doc.add_paragraph()

openings = [
    {
        "name": "Opening Version A - Contrarian Hook",
        "en": "Most people get this completely wrong.",
        "zh": "大多数人完全选错了。",
        "note": "Host looks directly at camera with a slight challenge expression. 2 seconds total.",
    },
    {
        "name": "Opening Version B - Curiosity Gap",
        "en": "Chinese medicine says the opposite of what you think.",
        "zh": "中医说的和你想的恰恰相反。",
        "note": "Host leans forward slightly. Mysterious tone. 2 seconds total.",
    },
    {
        "name": "Opening Version C - Challenge",
        "en": "9 out of 10 people choose the wrong answer. Are you the 1?",
        "zh": "10个人里9个选错。你是那1个吗？",
        "note": "Host holds up 3 fingers. Confident tone. 2-3 seconds total.",
    },
]

for o in openings:
    doc.add_heading(o["name"], level=2)
    p = doc.add_paragraph()
    p.add_run("English: ").bold = True
    p.add_run(o["en"])
    p2 = doc.add_paragraph()
    p2.add_run("Chinese: ").bold = True
    p2.add_run(o["zh"])
    p3 = doc.add_paragraph()
    p3.add_run("Animation Note: ").bold = True
    p3.add_run(o["note"])
    doc.add_paragraph()

doc.add_page_break()

doc.add_heading("FIXED ENDING ANIMATION (Same for all videos)", level=1)

endings = [
    {
        "name": "Ending Version A - Standard",
        "en": "Answer in comments. First correct answer wins a free body type quiz code!",
        "zh": "在评论区回答。第一个答对的送免费体质测试激活码！",
        "note": "Question image freezes. Text appears. 3 seconds total.",
    },
    {
        "name": "Ending Version B - Urgent",
        "en": "Comment your answer NOW. Correct answers get a free quiz code in your DM!",
        "zh": "立刻在评论区回答。答对的免费激活码发到你私信！",
        "note": "Question image freezes with countdown effect. 3 seconds total.",
    },
]

for e in endings:
    doc.add_heading(e["name"], level=2)
    p = doc.add_paragraph()
    p.add_run("English: ").bold = True
    p.add_run(e["en"])
    p2 = doc.add_paragraph()
    p2.add_run("Chinese: ").bold = True
    p2.add_run(e["zh"])
    p3 = doc.add_paragraph()
    p3.add_run("Animation Note: ").bold = True
    p3.add_run(e["note"])
    doc.add_paragraph()

doc.add_page_break()

doc.add_heading("VIDEO STRUCTURE (15 seconds total)", level=1)

structure = [
    ("0-2 sec", "Fixed Opening Animation", "Pre-made, same for all videos"),
    ("2-3 sec", "Question text appears", "Question image slide"),
    ("3-9 sec", "3 options shown (2 sec each)", "A -> B -> C, each on screen 2 seconds"),
    ("9-12 sec", "Options freeze", "All 3 options visible, frozen"),
    ("12-15 sec", "Fixed Ending Animation", "Pre-made, same for all videos"),
]

table = doc.add_table(rows=1, cols=3)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
hdr[0].text = "Time"
hdr[1].text = "Content"
hdr[2].text = "Note"
for time, content, note in structure:
    row = table.add_row().cells
    row[0].text = time
    row[1].text = content
    row[2].text = note

doc.add_paragraph()
doc.add_paragraph("Only the QUESTION IMAGE changes between videos. Opening and ending animations are fixed.", style="Intense Quote")

doc.add_page_break()

doc.add_heading("20 QUESTION IMAGES (Content for each video)", level=1)
doc.add_paragraph("Each video only needs ONE question image. The image contains the question + 3 options. Opening and ending are fixed animations.", style="Intense Quote")
doc.add_paragraph()

for q in questions:
    doc.add_heading(f"Episode {q['n']}", level=2)

    p = doc.add_paragraph()
    p.add_run("Question (English): ").bold = True
    p.add_run(q["q_en"])

    p = doc.add_paragraph()
    p.add_run("Question (Chinese): ").bold = True
    p.add_run(q["q_zh"])

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Options for Image:").bold = True

    doc.add_paragraph(q["a"], style="List Bullet")
    doc.add_paragraph(q["b"], style="List Bullet")
    doc.add_paragraph(q["c"], style="List Bullet")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Correct Answer: ").bold = True
    r = p.add_run(q["ans"])
    r.font.color.rgb = RGBColor(0xC9, 0xA3, 0x55)
    r.bold = True

    p = doc.add_paragraph()
    p.add_run("Why (English): ").bold = True
    p.add_run(q["why_en"])

    p = doc.add_paragraph()
    p.add_run("Why (Chinese): ").bold = True
    p.add_run(q["why_zh"])

    p = doc.add_paragraph()
    p.add_run("Comment Reply: ").bold = True
    p.add_run(f'Answer: {q["ans"]} ✓ {q["why_en"]} First correct answer wins a free body type quiz code! Check your DM 🎁')

    doc.add_paragraph("---")
    doc.add_paragraph()

output = os.path.join(os.getcwd(), "Short_Video_Scripts_Final_20.docx")
doc.save(output)
print(f"Saved: {output}")
