from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json, os

with open(os.path.join(os.path.dirname(__file__), "quiz-data.json"), "r", encoding="utf-8") as f:
    questions = json.load(f)

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# Title
title = doc.add_heading("Chinese Medicine Quiz Videos - Full Scripts", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("20 Videos x 15 Seconds Each\nFormat: Hook + Question + 3 Options + CTA")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0xC9, 0xA3, 0x55)

doc.add_paragraph()

# Structure guide
doc.add_heading("Video Structure (15 seconds)", level=1)
structure = [
    ("0-2s", "HOOK", "Digital human looks at camera and says the hook line"),
    ("2-3s", "QUESTION", "Question text appears on screen + voiceover"),
    ("3-9s", "OPTIONS", "3 options shown one by one (2s each)"),
    ("9-11s", "FREEZE", "All options visible + 'Answer in comments'"),
    ("11-13s", "REWARD", "'Correct answers get a free quiz code in DM'"),
    ("13-15s", "BRAND", "Logo + myeasterntype.com"),
]
for time, section, desc in structure:
    p = doc.add_paragraph()
    p.add_run(f"{time}  ").bold = True
    p.add_run(f"{section}: ").bold = True
    p.add_run(desc)

doc.add_page_break()

# Hook options
doc.add_heading("Hook Options (choose one for all videos)", level=1)
hooks = [
    ("Option A - Contrarian (Recommended)", '"Most people get this completely wrong." / 大多数人完全搞错了。'),
    ("Option B - Curiosity Gap", '"Chinese medicine says the opposite of what you think." / 中医说的和你想的恰恰相反。'),
    ("Option C - Challenge", '"9 out of 10 people choose the wrong answer. Are you the 1?" / 10个人里9个选错。你是那1个吗？'),
]
for name, text in hooks:
    p = doc.add_paragraph()
    p.add_run(f"{name}\n").bold = True
    p.add_run(text)
    doc.add_paragraph()

doc.add_page_break()

# Ending options
doc.add_heading("Ending Options (choose one for all videos)", level=1)
endings = [
    ("Option A - Standard (Recommended)", '"Answer in comments. First correct answer wins a free body type quiz code!" / 在评论区回答。第一个答对的赢免费体质测试码！'),
    ("Option B - Urgent", '"Comment your answer NOW. Correct answers get a free quiz code in your DM!" / 立刻在评论区回答。答对的在私信收到免费测试码！'),
]
for name, text in endings:
    p = doc.add_paragraph()
    p.add_run(f"{name}\n").bold = True
    p.add_run(text)
    doc.add_paragraph()

doc.add_page_break()

# Full scripts
doc.add_heading("Full Video Scripts (20 Videos)", level=1)

COMMENT_TEMPLATE_EN = "Answer: {ans}. {why_en} Most people pick the wrong one because it looks healthy by Western standards. First correct answer wins a free body type quiz code! Check your DM."
COMMENT_TEMPLATE_ZH = "答案：{ans}。{why_zh} 大多数人选错是因为按西方标准看起来健康。第一个答对的赢免费体质测试码！查收私信。"

for i, q in enumerate(questions, 1):
    doc.add_heading(f"Video {i}", level=2)

    # Hook
    p = doc.add_paragraph()
    p.add_run("[HOOK - 0-2s]\n").bold = True
    p.add_run('Digital human: "Most people get this completely wrong."\n')
    p.add_run("数字人：「大多数人完全搞错了。」")

    doc.add_paragraph()

    # Question
    p = doc.add_paragraph()
    p.add_run("[QUESTION - 2-3s]\n").bold = True
    p.add_run(f'{q["q_en"]}\n')
    p.add_run(f'{q["q_zh"]}')

    doc.add_paragraph()

    # Options
    p = doc.add_paragraph()
    p.add_run("[OPTIONS - 3-9s, 2s each]\n").bold = True
    doc.add_paragraph(q["a"])
    doc.add_paragraph(q["b"])
    doc.add_paragraph(q["c"])

    doc.add_paragraph()

    # Ending
    p = doc.add_paragraph()
    p.add_run("[ENDING - 9-15s]\n").bold = True
    p.add_run('Voiceover: "Answer in comments. First correct answer wins a free body type quiz code!"\n')
    p.add_run("画外音：「在评论区回答。第一个答对的赢免费体质测试码！」\n")
    p.add_run("[Brand logo + myeasterntype.com]")

    doc.add_paragraph()

    # Answer (for your reference, not in video)
    p = doc.add_paragraph()
    p.add_run("[CORRECT ANSWER - for your reference only]\n").bold = True
    p.add_run(f"Answer: {q['ans']}\n")
    p.add_run(f"EN: {q['why_en']}\n")
    p.add_run(f"ZH: {q['why_zh']}")

    doc.add_paragraph()

    # Comment template
    p = doc.add_paragraph()
    p.add_run("[COMMENT TO POST]\n").bold = True
    comment_en = COMMENT_TEMPLATE_EN.format(ans=q["ans"], why_en=q["why_en"])
    comment_zh = COMMENT_TEMPLATE_ZH.format(ans=q["ans"], why_zh=q["why_zh"])
    p.add_run(f"EN: {comment_en}\n\n")
    p.add_run(f"ZH: {comment_zh}")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("─" * 60)
    doc.add_paragraph()

output = os.path.join(os.path.dirname(__file__), "..", "Quiz_Video_Scripts_20.docx")
doc.save(output)
print(f"Saved: {output}")
print(f"Questions: {len(questions)}")
