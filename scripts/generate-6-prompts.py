from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

title = doc.add_heading("Opening & Ending Animation Prompts — 3 + 3", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("6 image prompts for image2. Pick 1 opening + 1 ending, use for all videos.")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0xC9, 0xA3, 0x55)

doc.add_paragraph()

doc.add_heading("Shared Visual Style (applies to all 6 images)", level=1)

style_items = [
    ("Resolution", "1000x1500px (2:3 vertical)"),
    ("Background", "Warm cream (#F5EFE6) with soft golden radial glow from center. Faint Chinese medicine herb illustrations (ginger, goji, chrysanthemum) at 5% opacity as watermark."),
    ("Host", "Young East Asian woman, late 20s, modern minimal olive-green or warm beige top. Natural makeup. Simple hair. Confident, approachable. Modern TCM practitioner vibe."),
    ("Colors", "Cream #F5EFE6, Gold #C9A355, Brown #2A1F14, Red accent #C75B3A"),
    ("Font", "Playfair Display serif for all text"),
    ("Footer", "Small 'myeasterntype.com' at bottom, 10pt, brown"),
    ("Vibe", "Premium wellness brand. Warm, modern, trustworthy. Not clinical. Not old-fashioned."),
]

for label, value in style_items:
    p = doc.add_paragraph()
    p.add_run(f"{label}: ").bold = True
    p.add_run(value)

doc.add_page_break()

# ===== 3 OPENINGS =====
doc.add_heading("3 OPENING ANIMATION PROMPTS", level=1)
doc.add_paragraph("Pick ONE. Use it for all 20 videos. The opening does NOT mention any question.", style="Intense Quote")

# Opening A
doc.add_heading("Opening A: Contrarian Hook", level=2)
p = doc.add_paragraph()
p.add_run("Text on image: ").bold = True
p.add_run("Most people get this completely wrong.")
p = doc.add_paragraph()
p.add_run("Host expression: ").bold = True
p.add_run("Confident, slight challenge, one eyebrow raised")
p = doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("image2 prompt (copy all below):").bold = True

doc.add_paragraph("""
Create a vertical image (1000x1500px, 2:3 ratio) for a short video opening frame.

Background: warm cream color (#F5EFE6) with a soft golden radial glow emanating from the center. Very faint illustrations of traditional Chinese medicine herbs (ginger root, goji berries, chrysanthemum flowers) as delicate watermark-style decorations at 5% opacity, barely visible.

Host: A young East Asian woman in her late 20s, wearing a modern minimal olive-green top. Natural makeup, simple hair style. She is looking directly at the camera with a confident, slightly challenging expression. One eyebrow slightly raised, as if saying "you probably don't know this." She is positioned on the right side of the frame, upper body visible.

Text overlay on the left side, in large elegant Playfair Display serif font, gold color (#C9A355):
"Most people get this completely wrong."

Below the text, a small decorative element: a thin gold line (1px) with a tiny diamond symbol in the center.

At the very bottom, small text: "myeasterntype.com" in brown (#2A1F14), 10pt size.

Style: premium wellness brand, warm and modern, clean composition with generous whitespace. Professional but approachable. NOT clinical, NOT old-fashioned.""")

doc.add_paragraph()
doc.add_paragraph("Animation note: Host holds still 1.5s, slight head tilt at 'wrong.' Total 2s.")
doc.add_paragraph()

# Opening B
doc.add_heading("Opening B: Curiosity Gap", level=2)
p = doc.add_paragraph()
p.add_run("Text on image: ").bold = True
p.add_run("Chinese medicine says the opposite of what you think.")
p = doc.add_paragraph()
p.add_run("Host expression: ").bold = True
p.add_run("Mysterious knowing smile, leaning forward slightly")
p = doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("image2 prompt (copy all below):").bold = True

doc.add_paragraph("""
Create a vertical image (1000x1500px, 2:3 ratio) for a short video opening frame.

Background: warm cream color (#F5EFE6) with a soft golden radial glow from center. Very faint illustrations of Chinese medicine herbs (lotus seed pod, red dates, yam root) as delicate watermark decorations at 5% opacity.

Host: A young East Asian woman in her late 20s, wearing a modern minimal warm beige top. Natural makeup, simple hair. She is leaning forward slightly toward the camera with a mysterious, knowing smile. Eyes slightly narrowed as if sharing a secret. Positioned on the right side, upper body visible.

Text overlay on the left side, in large elegant Playfair Display serif font, gold color (#C9A355):
"Chinese medicine says the opposite of what you think."

The word "opposite" is highlighted in a slightly larger size or warmer gold tone to emphasize the twist.

Below the text, a small decorative element: a thin gold line with a tiny lotus symbol in the center.

At the very bottom, small text: "myeasterntype.com" in brown (#2A1F14), 10pt size.

Style: premium wellness brand, warm and intriguing, clean composition. The host expression should convey "I know something you don't." Modern Chinese medicine meets lifestyle brand aesthetic.""")

doc.add_paragraph()
doc.add_paragraph("Animation note: Host leans in slowly, smile appears at 'opposite.' Total 2-3s.")
doc.add_paragraph()

# Opening C
doc.add_heading("Opening C: Challenge", level=2)
p = doc.add_paragraph()
p.add_run("Text on image: ").bold = True
p.add_run("9 out of 10 people choose the wrong answer. Are you the 1?")
p = doc.add_paragraph()
p.add_run("Host expression: ").bold = True
p.add_run("Confident, playful, holding up 3 fingers, game-show energy")
p = doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("image2 prompt (copy all below):").bold = True

doc.add_paragraph("""
Create a vertical image (1000x1500px, 2:3 ratio) for a short video opening frame.

Background: warm cream color (#F5EFE6) with a soft golden radial glow from center. Very faint illustrations of Chinese medicine herbs (walnut, black sesame, dried tangerine peel) as delicate watermark decorations at 5% opacity.

Host: A young East Asian woman in her late 20s, wearing a modern minimal deep teal or forest green top. Natural makeup, simple hair. She is holding up three fingers toward the camera in a confident, playful gesture. Expression is challenging and inviting, like a game show host who knows the answer. Direct eye contact. Positioned center-right, upper body visible.

Text overlay on the upper-left area, in large elegant Playfair Display serif font, gold color (#C9A355):
"9 out of 10 people choose the wrong answer."

Below that line, slightly smaller text:
"Are you the 1?"

The number "1" is highlighted in a warm red accent (#C75B3A) to make it pop.

Below the text, a small decorative element: a thin gold line with a tiny diamond symbol.

At the very bottom, small text: "myeasterntype.com" in brown (#2A1F14), 10pt size.

Style: premium wellness brand, warm and engaging, gamified feel. The host should look like she is inviting you to a challenge.""")

doc.add_paragraph()
doc.add_paragraph("Animation note: Host raises 3 fingers on '9 out of 10', points to camera on 'Are you the 1.' Total 2-3s.")

doc.add_page_break()

# ===== 3 ENDINGS =====
doc.add_heading("3 ENDING ANIMATION PROMPTS", level=1)
doc.add_paragraph("Pick ONE. Use it for all 20 videos. The ending directs users to comment.", style="Intense Quote")

# Ending A
doc.add_heading("Ending A: Standard", level=2)
p = doc.add_paragraph()
p.add_run("Text on image: ").bold = True
p.add_run("Answer in comments. First correct answer wins a free body type quiz code!")
p = doc.add_paragraph()
p.add_run("Host action: ").bold = True
p.add_run("Warm smile, gesturing toward comment section")
p = doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("image2 prompt (copy all below):").bold = True

doc.add_paragraph("""
Create a vertical image (1000x1500px, 2:3 ratio) for a short video ending frame.

Background: warm cream color (#F5EFE6) with a soft golden radial glow from center. Very faint Chinese medicine herb illustrations at 5% opacity, same style as the opening frame.

Host: Same young East Asian woman in her late 20s, wearing the same modern minimal olive-green top. She now has a warm, welcoming smile. Both hands gesturing downward toward where the comment section would be (pointing down or making a typing gesture). Expression: encouraging and inviting.

Text overlay, centered, in Playfair Display serif font:
Line 1 (gold, #C9A355, large): "Answer in comments."
Line 2 (brown, #2A1F14, medium): "First correct answer wins a free"
Line 3 (gold, #C9A355, medium, bold): "body type quiz code!"

A small gift box icon next to line 3.

At the very bottom, small text: "myeasterntype.com" in brown (#2A1F14), 10pt size.

Style: same premium wellness aesthetic as the opening. Clean, warm, inviting.""")

doc.add_paragraph()
doc.add_paragraph("Animation note: Host holds smile, points down 1.5s. Text fades in. Total 3s.")
doc.add_paragraph()

# Ending B
doc.add_heading("Ending B: Urgent", level=2)
p = doc.add_paragraph()
p.add_run("Text on image: ").bold = True
p.add_run("Comment your answer NOW. Correct answers get a free quiz code in your DM!")
p = doc.add_paragraph()
p.add_run("Host action: ").bold = True
p.add_run("Excited, one finger pointing at camera, urgent energy")
p = doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("image2 prompt (copy all below):").bold = True

doc.add_paragraph("""
Create a vertical image (1000x1500px, 2:3 ratio) for a short video ending frame.

Background: warm cream color (#F5EFE6) with a soft golden radial glow from center. Same faint herb illustrations as opening, 5% opacity.

Host: Same young East Asian woman in her late 20s, same olive-green top. She now has an excited, urgent expression. One hand pointing directly at the camera, other hand holding a small gift box icon or envelope icon (representing the DM code). Expression: "hurry up and comment!"

Text overlay, centered, in Playfair Display serif font:
Line 1 (gold, #C9A355, large, bold): "Comment NOW!"
Line 2 (brown, #2A1F14, medium): "Correct answers get a free"
Line 3 (gold, #C9A355, medium): "quiz code in your DM!"

A small gift icon next to the text.

At the very bottom, small text: "myeasterntype.com" in brown (#2A1F14), 10pt size.

Style: same premium wellness aesthetic but with more energy and urgency. Clean, warm, exciting.""")

doc.add_paragraph()
doc.add_paragraph("Animation note: Host points at camera, text pops in with slight bounce. Total 3s.")
doc.add_paragraph()

# Ending C
doc.add_heading("Ending C: Teaser", level=2)
p = doc.add_paragraph()
p.add_run("Text on image: ").bold = True
p.add_run("Think you got it right? Comment your answer. Free quiz code for correct answers!")
p = doc.add_paragraph()
p.add_run("Host action: ").bold = True
p.add_run("Head tilted, curious, finger on chin, playful doubt")
p = doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("image2 prompt (copy all below):").bold = True

doc.add_paragraph("""
Create a vertical image (1000x1500px, 2:3 ratio) for a short video ending frame.

Background: warm cream color (#F5EFE6) with a soft golden radial glow from center. Same faint herb illustrations as opening, 5% opacity.

Host: Same young East Asian woman in her late 20s, same olive-green top. Head tilted slightly to one side, one finger resting on her chin. Expression: playful doubt, like "are you sure you got it right?" A slight smirk. Curious and teasing.

Text overlay, centered, in Playfair Display serif font:
Line 1 (gold, #C9A355, large): "Think you got it right?"
Line 2 (brown, #2A1F14, medium): "Comment your answer."
Line 3 (gold, #C9A355, medium, bold): "Free quiz code for correct answers!"

A small gift icon at the end of line 3.

At the very bottom, small text: "myeasterntype.com" in brown (#2A1F14), 10pt size.

Style: same premium wellness aesthetic. Playful, teasing, engaging. Makes the viewer want to prove they know the answer.""")

doc.add_paragraph()
doc.add_paragraph("Animation note: Host tilts head, finger on chin. Text fades in. Total 3s.")

doc.add_page_break()

doc.add_heading("How to Use", level=1)

steps = [
    "Pick 1 opening (A, B, or C) and copy the prompt into image2",
    "Pick 1 ending (A, B, or C) and copy the prompt into image2",
    "Make 2 images total. These are your fixed opening and ending frames",
    "Turn each image into a 2-3 second animation (you can add motion in your video editor)",
    "For each of the 20 videos: opening animation + question image + ending animation = 12-15 seconds",
    "Opening and ending stay the same for ALL 20 videos. Only the middle question image changes",
]

for i, step in enumerate(steps, 1):
    p = doc.add_paragraph(style="List Number")
    p.add_run(step)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Tip: ").bold = True
p.add_run("Keep the host outfit, background, and colors IDENTICAL across opening and ending. This builds brand recognition. The host should look like the same person in the same setting.")

output = os.path.join(os.getcwd(), "Opening_Ending_6_Prompts.docx")
doc.save(output)
print(f"Saved: {output}")
